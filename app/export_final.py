# -*- coding: utf-8 -*-
"""
SafeFill - 最终文件导出脚本
功能：读取 review_results 中用户确认后的检查结果，
     将其应用到草稿文件副本上，生成最终版文件。
安全约束：不修改原始文件、不修改草稿文件、不修改 vault。
"""

import os
import sys
import json
import re
import shutil
from datetime import datetime
from pathlib import Path

# ------------------------------------------------------------
# 项目根目录
# ------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent

REVIEW_RESULTS_DIR = PROJECT_ROOT / "review_results"
DRAFT_DIR = PROJECT_ROOT / "draft_outputs"
FINAL_OUTPUT_DIR = PROJECT_ROOT / "final_outputs"
FINAL_REPORTS_DIR = PROJECT_ROOT / "final_reports"
LOGS_DIR = PROJECT_ROOT / "logs"

# ------------------------------------------------------------
# 字段名映射（复用第 2 阶段格式）
# ------------------------------------------------------------
FIELD_NAME_MAP = {
    "姓名": "name", "名字": "name", "申报人": "name", "申请人": "name", "负责人": "name",
    "性别": "gender",
    "出生日期": "birth_date", "出生年月": "birth_date", "生日": "birth_date", "出生年月日": "birth_date",
    "身份证号": "id_number", "身份证号码": "id_number", "证件号码": "id_number",
    "身份证": "id_number", "公民身份号码": "id_number",
    "手机号": "phone", "手机号码": "phone", "联系电话": "phone", "电话": "phone",
    "联系方式": "phone", "移动电话": "phone",
    "邮箱": "email", "电子邮件": "email", "电子邮箱": "email",
    "E-mail": "email", "Email": "email", "邮件": "email",
    "工作单位": "organization", "单位": "organization", "所在单位": "organization",
    "依托单位": "organization", "申报单位": "organization", "单位名称": "organization",
    "部门": "department", "院系": "department", "学院": "department",
    "所在部门": "department", "所在院系": "department", "系所": "department",
    "职称": "title", "专业技术职务": "title", "现任职称": "title", "职务": "title",
    "学历": "education", "最高学历": "education", "文化程度": "education",
    "学位": "degree", "最高学位": "degree", "学位名称": "degree",
    "专业": "major", "所学专业": "major", "攻读专业": "major",
    "专业方向": "research_area", "研究方向": "research_area", "研究领域": "research_area",
    "学科方向": "research_area", "从事专业": "research_area",
    "通讯地址": "address", "通信地址": "address", "联系地址": "address",
    "地址": "address", "家庭地址": "address", "住址": "address", "邮寄地址": "address",
    "个人简介": "biography", "个人简历": "biography", "简介": "biography",
    "简历": "biography", "个人陈述": "biography", "自我简介": "biography",
    "项目经历": "project_experience", "科研项目": "project_experience",
    "主持项目": "project_experience", "参与项目": "project_experience",
    "项目列表": "project_experience", "承担项目": "project_experience",
}

SENSITIVE_FIELDS = {"id_number", "phone", "address", "photo_path"}
LONG_TEXT_FIELDS = {"biography", "project_experience"}


# ------------------------------------------------------------
# 日志
# ------------------------------------------------------------
def write_log(message: str):
    LOGS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_file = LOGS_DIR / "export_final.log"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {message}\n")
    print(f"[LOG] {message}")


# ------------------------------------------------------------
# 脱敏
# ------------------------------------------------------------
def mask_id_number(value: str) -> str:
    v = str(value).strip()
    if len(v) >= 15: return v[:3] + "*" * (len(v) - 7) + v[-4:]
    if len(v) >= 8: return v[:2] + "*" * (len(v) - 4) + v[-2:]
    return "***"

def mask_phone(value: str) -> str:
    v = str(value).strip().replace("-", "").replace(" ", "")
    if len(v) == 11 and v.isdigit(): return v[:3] + "****" + v[-4:]
    if len(v) >= 8: return v[:3] + "****" + v[-2:]
    return "***"

def mask_address(value: str) -> str:
    v = str(value).strip()
    if len(v) <= 6: return v + "***"
    return v[:6] + "***"

def mask_for_display(fkey: str, value) -> str:
    if not value: return ""
    s = str(value)
    if fkey == "id_number": return mask_id_number(s)
    if fkey == "phone": return mask_phone(s)
    if fkey == "address": return mask_address(s)
    if fkey == "photo_path": return "存在" if s.strip() else "不存在"
    return s


# ------------------------------------------------------------
# review_result 查找
# ------------------------------------------------------------
def _is_test_file(name: str) -> bool:
    """排除测试/示例 review_result。"""
    lower = name.lower()
    for kw in ("_test_", "html_test", "sample", "demo"):
        if kw in lower:
            return True
    return False

def find_latest_review() -> Path | None:
    """找到最新的正式 review_result_*.json（按修改时间，排除测试文件）。"""
    if not REVIEW_RESULTS_DIR.exists():
        return None
    files = [f for f in REVIEW_RESULTS_DIR.glob("review_result_*.json")
             if not _is_test_file(f.name)]
    if not files:
        return None
    files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
    return files[0]


# ------------------------------------------------------------
# review_result 有效性检查
# ------------------------------------------------------------
def validate_review_freshness(review_path: Path, review: dict) -> bool:
    """
    检查 review_result 是否与本轮 FormReview 生成的一致。
    返回 True = 通过检查。
    """
    latest_html_path = PROJECT_ROOT / "review_html" / "latest_review_html.json"
    if not latest_html_path.exists():
        print("[WARN] 未发现 latest_review_html.json，将按旧流程导出。")
        return True

    try:
        with open(latest_html_path, "r", encoding="utf-8") as f:
            latest = json.load(f)
    except Exception:
        print("[WARN] latest_review_html.json 读取失败，将按旧流程导出。")
        return True

    valid_drafts = {item.get("draft_file", "") for item in latest.get("items", [])}
    review_draft = review.get("draft_file", "")

    if review_draft not in valid_drafts:
        print(f"\n[STOP] 当前 review_result 不是本轮 FormReview 生成的检查结果。")
        print(f"  review_result draft: {review_draft}")
        print(f"  本轮 draft 列表: {valid_drafts}")
        print(f"  请回到网页 http://127.0.0.1:8787/ 点击【保存检查结果】后再导出。\n")
        write_log("SECURITY: review_result 与 latest_review_html 不一致，拒绝导出")
        return False

    # html_review 模式额外校验
    if review.get("mode") == "html_review":
        if not review.get("confirmed_by_user"):
            print(f"\n[STOP] review_result 中 confirmed_by_user 不为 true。")
            write_log("SECURITY: confirmed_by_user 不为 true，拒绝导出")
            return False
        if not review.get("tables"):
            print(f"\n[STOP] review_result 中 tables 为空。")
            write_log("SECURITY: tables 为空，拒绝导出")
            return False

    return True


def show_export_summary(review_path: Path, review: dict, draft_path: Path, ext: str):
    """打印导出前摘要。"""
    mode = review.get("mode", "(旧版字段确认模式)")
    source = review.get("source_file", "")
    print(f"\n{'='*55}")
    print(f"  导出前摘要")
    print(f"  {'-'*45}")
    print(f"  review_result: {review_path.name}")
    print(f"  创建时间:      {review.get('created_at', '未知')}")
    print(f"  mode:          {mode}")
    print(f"  source_file:   {Path(source).name if source else '未知'}")
    print(f"  draft_file:    {draft_path.name}")
    print(f"  输出目录:      {FINAL_OUTPUT_DIR}")
    print(f"  {'-'*45}")
    print(f"  安全确认:")
    print(f"    修改原始文件: 否")
    print(f"    修改草稿文件: 否")
    print(f"    修改 vault:   否")
    print(f"    联网:         否")
    print(f"    调用 API:     否")
    print(f"{'='*55}")


def validate_draft_path(draft_file_str: str) -> Path:
    """Return a draft path only when it resolves inside draft_outputs."""
    if not draft_file_str:
        raise ValueError("review_result 中未记录 draft_file 路径")

    draft_path = Path(draft_file_str)
    if not draft_path.is_absolute():
        draft_path = DRAFT_DIR / draft_path.name

    draft_resolved = draft_path.resolve()
    try:
        draft_resolved.relative_to(DRAFT_DIR.resolve())
    except ValueError as e:
        raise PermissionError(f"draft_file 指向 draft_outputs 目录之外: {draft_file_str}") from e
    if not draft_resolved.exists():
        raise FileNotFoundError(f"草稿文件不存在: {draft_resolved}")
    return draft_resolved


def make_final_path(draft_path: Path, timestamp: str) -> Path:
    FINAL_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    final_name = f"{draft_path.stem}_最终版{draft_path.suffix.lower()}"
    final_path = FINAL_OUTPUT_DIR / final_name
    if final_path.exists():
        final_name = f"{draft_path.stem}_最终版_{timestamp}{draft_path.suffix.lower()}"
        final_path = FINAL_OUTPUT_DIR / final_name
    return final_path


# ------------------------------------------------------------
# 字段名匹配
# ------------------------------------------------------------
def match_field_name(cell_text: str) -> str | None:
    text = cell_text.strip()
    if not text: return None
    if text in FIELD_NAME_MAP: return FIELD_NAME_MAP[text]
    cleaned = re.sub(r'\s*[：:]\s*$', '', text).strip()
    if cleaned in FIELD_NAME_MAP: return FIELD_NAME_MAP[cleaned]
    for cn, fk in FIELD_NAME_MAP.items():
        if cn in text and len(cn) >= 2: return fk
    return None

def is_empty_cell(value) -> bool:
    return value is None or (isinstance(value, str) and value.strip() == "")


# ------------------------------------------------------------
# 查找字段值（遍历待处理列表）
# ------------------------------------------------------------
def find_field_value(field_label: str, field_key: str, collection: list) -> str | None:
    """在 collection 中查找匹配的字段值。"""
    for item in collection:
        if not isinstance(item, dict): continue
        ik = item.get("field_key", "")
        il = item.get("field_label", "")
        if ik == field_key or il == field_label or field_label in il or il in field_label:
            val = item.get("new_value") or item.get("value")
            if val: return str(val)
    return None


def is_field_rejected(field_label: str, field_key: str, rejected: list) -> bool:
    for item in rejected:
        if not isinstance(item, dict): continue
        if item.get("field_key") == field_key or item.get("field_label") == field_label:
            return True
    return False


# ------------------------------------------------------------
# docx 导出
# ------------------------------------------------------------
def export_docx(draft_path: Path, final_path: Path, review: dict) -> dict:
    import docx
    report = {
        "written": [], "skipped": [], "rejected_cleared": [],
        "not_located": [], "long_text_warning": [], "not_supported": [],
    }
    try:
        doc = docx.Document(str(draft_path))
    except Exception as e:
        report["not_supported"].append(f"无法打开草稿: {e}")
        return report

    # 收集待处理字段
    confirmed = review.get("confirmed_fields", [])
    modified = review.get("modified_fields", [])
    filled_by_user = review.get("unknown_fields_filled_by_user", [])
    rejected = review.get("fields_rejected", [])
    all_actions = confirmed + modified + filled_by_user

    cells_filled = 0
    cells_cleared = 0

    for table_idx, table in enumerate(doc.tables):
        rows = table.rows
        for row_idx, row in enumerate(rows):
            cells = row.cells
            for col_idx in range(len(cells)):
                cell = cells[col_idx]
                cell_text = cell.text.strip()
                fkey = match_field_name(cell_text)
                if fkey is None: continue

                target_col = col_idx + 1
                if target_col >= len(cells): continue
                target_cell = cells[target_col]

                loc_desc = f"Table[{table_idx+1}] R{row_idx+1}"

                # 处理 rejected
                if is_field_rejected(cell_text, fkey, rejected):
                    if not is_empty_cell(target_cell.text):
                        try:
                            for para in target_cell.paragraphs:
                                for run in para.runs:
                                    run.text = ""
                            cells_cleared += 1
                            report["rejected_cleared"].append({
                                "field_label": cell_text, "field_key": fkey, "location": loc_desc,
                            })
                            write_log(f"INFO: 已清空被拒字段 {cell_text}")
                        except Exception as e:
                            report["skipped"].append({
                                "field_label": cell_text, "field_key": fkey,
                                "reason": f"清空失败: {e}", "location": loc_desc,
                            })
                    else:
                        report["rejected_cleared"].append({
                            "field_label": cell_text, "field_key": fkey, "location": loc_desc,
                            "note": "本身为空",
                        })
                    continue

                # 查找值：优先 modified/filled_by_user，然后 confirmed
                value = find_field_value(cell_text, fkey, modified) or \
                        find_field_value(cell_text, fkey, filled_by_user) or \
                        find_field_value(cell_text, fkey, confirmed)
                if value is None: continue

                target_text = target_cell.text.strip()

                # 决定是否覆盖
                is_from_mod = find_field_value(cell_text, fkey, modified) is not None
                is_from_user = find_field_value(cell_text, fkey, filled_by_user) is not None

                if not is_empty_cell(target_text) and not is_from_mod:
                    # 已有内容且不是用户明确修改，不覆盖
                    report["skipped"].append({
                        "field_label": cell_text, "field_key": fkey,
                        "reason": "右侧单元格已有内容，非用户修改不覆盖",
                        "location": loc_desc, "existing": target_text[:40],
                    })
                    continue

                try:
                    for para in target_cell.paragraphs:
                        for run in para.runs:
                            run.text = ""
                    if target_cell.paragraphs:
                        target_cell.paragraphs[0].add_run(value)
                    else:
                        target_cell.add_paragraph(value)
                    cells_filled += 1

                    entry = {
                        "field_label": cell_text, "field_key": fkey,
                        "value_length": len(value), "location": loc_desc,
                    }
                    report["written"].append(entry)
                    if fkey in LONG_TEXT_FIELDS:
                        report["long_text_warning"].append(entry)
                    write_log(f"INFO: 已写入 {cell_text} (字段: {fkey})")
                except Exception as e:
                    report["skipped"].append({
                        "field_label": cell_text, "field_key": fkey,
                        "reason": f"写入失败: {e}", "location": loc_desc,
                    })

    # 检查无法定位的字段
    all_labels_written = {e["field_label"] for e in report["written"]}
    for item in all_actions:
        label = item.get("field_label", "")
        if label and label not in all_labels_written:
            report["not_located"].append({
                "field_label": label, "field_key": item.get("field_key", ""),
                "reason": "在表格中未找到对应字段名",
            })

    report["_stats"] = {"cells_filled": cells_filled, "cells_cleared": cells_cleared}

    # 保存
    doc.save(str(final_path))
    return report


# ------------------------------------------------------------
# xlsx 导出
# ------------------------------------------------------------
def export_xlsx(draft_path: Path, final_path: Path, review: dict) -> dict:
    import openpyxl
    report = {
        "written": [], "skipped": [], "rejected_cleared": [],
        "not_located": [], "long_text_warning": [], "not_supported": [],
    }
    try:
        wb = openpyxl.load_workbook(str(draft_path))
    except Exception as e:
        report["not_supported"].append(f"无法打开草稿: {e}")
        return report

    confirmed = review.get("confirmed_fields", [])
    modified = review.get("modified_fields", [])
    filled_by_user = review.get("unknown_fields_filled_by_user", [])
    rejected = review.get("fields_rejected", [])
    all_actions = confirmed + modified + filled_by_user

    cells_filled = 0
    cells_cleared = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows():
            row_values = [c.value for c in row]
            for col_idx in range(len(row)):
                cell = row[col_idx]
                cv = cell.value
                if cv is None: continue
                cell_text = str(cv).strip()
                if not cell_text: continue

                fkey = match_field_name(cell_text)
                if fkey is None: continue

                target_col = col_idx + 1
                if target_col >= len(row): continue
                target_cell = row[target_col]
                loc_desc = f"Sheet[{sheet_name}] R{cell.row}C{target_cell.column}"

                # 合并单元格检查
                is_merged = False
                for mr in ws.merged_cells.ranges:
                    if target_cell.coordinate in mr:
                        if target_cell.column != mr.min_col or target_cell.row != mr.min_row:
                            report["skipped"].append({
                                "field_label": cell_text, "field_key": fkey,
                                "reason": f"目标在合并区域 {mr} 非左上角", "location": loc_desc,
                            })
                            is_merged = True
                            break
                if is_merged: continue

                # rejected
                if is_field_rejected(cell_text, fkey, rejected):
                    if not is_empty_cell(target_cell.value):
                        target_cell.value = None
                        cells_cleared += 1
                        report["rejected_cleared"].append({
                            "field_label": cell_text, "field_key": fkey, "location": loc_desc,
                        })
                    else:
                        report["rejected_cleared"].append({
                            "field_label": cell_text, "field_key": fkey,
                            "location": loc_desc, "note": "本身为空",
                        })
                    continue

                # 找值
                value = find_field_value(cell_text, fkey, modified) or \
                        find_field_value(cell_text, fkey, filled_by_user) or \
                        find_field_value(cell_text, fkey, confirmed)
                if value is None: continue

                is_from_mod = find_field_value(cell_text, fkey, modified) is not None
                if not is_empty_cell(target_cell.value) and not is_from_mod:
                    report["skipped"].append({
                        "field_label": cell_text, "field_key": fkey,
                        "reason": "右侧已有内容，非用户修改不覆盖",
                        "location": loc_desc,
                        "existing": str(target_cell.value)[:40],
                    })
                    continue

                try:
                    target_cell.value = value
                    cells_filled += 1
                    entry = {
                        "field_label": cell_text, "field_key": fkey,
                        "value_length": len(value), "location": loc_desc,
                    }
                    report["written"].append(entry)
                    if fkey in LONG_TEXT_FIELDS:
                        report["long_text_warning"].append(entry)
                    write_log(f"INFO: 已写入 {cell_text} (字段: {fkey})")
                except Exception as e:
                    report["skipped"].append({
                        "field_label": cell_text, "field_key": fkey,
                        "reason": f"写入失败: {e}", "location": loc_desc,
                    })

    # 检查无法定位
    all_written = {e["field_label"] for e in report["written"]}
    for item in all_actions:
        label = item.get("field_label", "")
        if label and label not in all_written:
            report["not_located"].append({
                "field_label": label, "field_key": item.get("field_key", ""),
                "reason": "在表格中未找到对应字段名",
            })

    report["_stats"] = {"cells_filled": cells_filled, "cells_cleared": cells_cleared}

    wb.save(str(final_path))
    return report


# ------------------------------------------------------------
# html_review 导出
# ------------------------------------------------------------
def _set_docx_cell_text(cell, value: str):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(value)
    else:
        cell.add_paragraph(value)


def export_html_review_docx(final_path: Path, review: dict) -> dict:
    import docx

    report = {
        "written": [], "skipped": [], "rejected_cleared": [],
        "not_located": [], "long_text_warning": [], "not_supported": [],
        "mode": "html_review",
    }
    try:
        doc = docx.Document(str(final_path))
    except Exception as e:
        report["not_supported"].append(f"无法打开最终 Word 文件: {e}")
        return report

    cells_written = 0
    cells_skipped = 0
    for table_data in review.get("tables", []):
        if not isinstance(table_data, dict):
            continue
        table_index = table_data.get("table_index", 1)
        try:
            doc_table_index = int(table_index) - 1
        except (TypeError, ValueError):
            doc_table_index = 0
        rows_data = table_data.get("rows", [])

        if doc_table_index < 0 or doc_table_index >= len(doc.tables):
            cells_skipped += sum(len(r) for r in rows_data if isinstance(r, list))
            report["skipped"].append({
                "location": f"Table[{table_index}]",
                "reason": "review_result 中的表格索引超过 Word 表格数量",
            })
            continue

        table = doc.tables[doc_table_index]
        for row_idx, row_values in enumerate(rows_data):
            if not isinstance(row_values, list):
                continue
            if row_idx >= len(table.rows):
                cells_skipped += len(row_values)
                report["skipped"].append({
                    "location": f"Table[{table_index}] R{row_idx + 1}",
                    "reason": "review_result 行数超过 Word 表格行数",
                })
                continue

            cells = table.rows[row_idx].cells
            for col_idx, value in enumerate(row_values):
                if col_idx >= len(cells):
                    cells_skipped += 1
                    report["skipped"].append({
                        "location": f"Table[{table_index}] R{row_idx + 1}C{col_idx + 1}",
                        "reason": "review_result 列数超过 Word 表格列数",
                    })
                    continue

                text = "" if value is None else str(value)
                try:
                    _set_docx_cell_text(cells[col_idx], text)
                    cells_written += 1
                    report["written"].append({
                        "field_label": f"cell R{row_idx + 1}C{col_idx + 1}",
                        "field_key": "html_cell",
                        "value": text,
                        "value_length": len(text),
                        "location": f"Table[{table_index}] R{row_idx + 1}C{col_idx + 1}",
                    })
                except Exception as e:
                    cells_skipped += 1
                    report["skipped"].append({
                        "location": f"Table[{table_index}] R{row_idx + 1}C{col_idx + 1}",
                        "reason": f"写入 Word 单元格失败: {e}",
                    })

    report["_stats"] = {"cells_written": cells_written, "cells_skipped": cells_skipped}
    doc.save(str(final_path))
    return report


def _is_non_top_left_merged_cell(ws, cell) -> bool:
    for merged_range in ws.merged_cells.ranges:
        if cell.coordinate in merged_range:
            return cell.row != merged_range.min_row or cell.column != merged_range.min_col
    return False


def export_html_review_xlsx(final_path: Path, review: dict) -> dict:
    import openpyxl

    report = {
        "written": [], "skipped": [], "rejected_cleared": [],
        "not_located": [], "long_text_warning": [], "not_supported": [],
        "mode": "html_review",
    }
    try:
        wb = openpyxl.load_workbook(str(final_path))
    except Exception as e:
        report["not_supported"].append(f"无法打开最终 Excel 文件: {e}")
        return report

    cells_written = 0
    cells_skipped = 0
    worksheets = wb.worksheets
    for table_data in review.get("tables", []):
        if not isinstance(table_data, dict):
            continue
        table_index = table_data.get("table_index", 1)
        try:
            sheet_index = int(table_index) - 1
        except (TypeError, ValueError):
            sheet_index = 0
        if sheet_index < 0 or sheet_index >= len(worksheets):
            sheet_index = 0
        ws = worksheets[sheet_index]

        for row_idx, row_values in enumerate(table_data.get("rows", []), start=1):
            if not isinstance(row_values, list):
                continue
            for col_idx, value in enumerate(row_values, start=1):
                cell = ws.cell(row=row_idx, column=col_idx)
                location = f"Sheet[{ws.title}] R{row_idx}C{col_idx}"
                if _is_non_top_left_merged_cell(ws, cell):
                    cells_skipped += 1
                    report["skipped"].append({
                        "location": location,
                        "reason": "目标单元格位于合并区域且不是左上角",
                    })
                    continue

                text = "" if value is None else str(value)
                try:
                    cell.value = text
                    cells_written += 1
                    report["written"].append({
                        "field_label": f"cell R{row_idx}C{col_idx}",
                        "field_key": "html_cell",
                        "value": text,
                        "value_length": len(text),
                        "location": location,
                    })
                except Exception as e:
                    cells_skipped += 1
                    report["skipped"].append({
                        "location": location,
                        "reason": f"写入 Excel 单元格失败: {e}",
                    })

    report["_stats"] = {"cells_written": cells_written, "cells_skipped": cells_skipped}
    wb.save(str(final_path))
    return report


# ------------------------------------------------------------
# 报告生成
# ------------------------------------------------------------
def generate_final_reports(review_path: Path, draft_path: Path, final_path: Path,
                           original_file: str, report: dict, review: dict, timestamp: str):
    FINAL_REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    # ---- JSON ----
    json_data = {
        "_说明": "最终导出报告 - 仅供本地使用",
        "_生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "mode": review.get("mode", "legacy"),
        "review_result文件": str(review_path),
        "草稿文件": str(draft_path),
        "原始文件": original_file,
        "最终文件": str(final_path),
        "成功写入字段": report.get("written", []),
        "跳过字段": report.get("skipped", []),
        "用户拒绝并清空": report.get("rejected_cleared", []),
        "用户修改字段": review.get("modified_fields", []),
        "用户补充字段": review.get("unknown_fields_filled_by_user", []),
        "建议保存到资料库": review.get("fields_marked_for_profile_save", []),
        "无法定位字段": report.get("not_located", []),
        "长文本复查提醒": report.get("long_text_warning", []),
        "安全确认": {
            "修改原始文件": False, "修改草稿文件": False,
            "修改vault": False, "联网": False, "调用API": False,
        },
        "_stats": report.get("_stats", {}),
    }
    jp = FINAL_REPORTS_DIR / f"final_report_{timestamp}.json"
    with open(jp, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)

    # ---- Markdown ----
    md = [
        "# 📄 最终导出报告",
        f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"> mode：{review.get('mode', 'legacy')}",
        f"> review_result：{review_path.name}",
        f"> 草稿文件：{draft_path.name}",
        f"> 最终文件：{final_path.name}",
        f"> 原始文件：{original_file}",
        "", "---", "",
        "## 🔒 安全确认",
        "| 项目 | 状态 |",
        "|------|------|",
        "| 修改原始文件 | ❌ 否 |",
        "| 修改草稿文件 | ❌ 否 |",
        "| 修改 vault | ❌ 否 |",
        "| 联网 | ❌ 否 |",
        "| 调用 API | ❌ 否 |",
        "", "---", "",
    ]

    written = report.get("written", [])
    if written:
        md.append("## ✅ 成功写入字段")
        md.append("| 字段名 | 资料库字段 | 写入值（脱敏） | 位置 | 备注 |")
        md.append("|--------|-----------|---------------|------|------|")
        for e in written:
            fk = e.get("field_key", "")
            display = mask_for_display(fk, e.get("value", ""))
            if len(display) > 40: display = display[:40] + "..."
            display = display.replace("|", "/")
            note = "⚠️ 长文本需复查" if fk in LONG_TEXT_FIELDS else ""
            md.append(f"| {e.get('field_label','?')} | {fk} | {display} | {e.get('location','')} | {note} |")
        md.append("")

    skipped = report.get("skipped", [])
    if skipped:
        md.append("## ⚠️ 跳过字段")
        md.append("| 字段名 | 原因 | 位置 |")
        md.append("|--------|------|------|")
        for e in skipped:
            md.append(f"| {e.get('field_label','?')} | {e.get('reason','')} | {e.get('location','')} |")
        md.append("")

    rejected = report.get("rejected_cleared", [])
    if rejected:
        md.append("## ❌ 用户拒绝字段")
        md.append("| 字段名 | 位置 | 备注 |")
        md.append("|--------|------|------|")
        for e in rejected:
            md.append(f"| {e.get('field_label','?')} | {e.get('location','')} | {e.get('note','已清空')} |")
        md.append("")

    not_located = report.get("not_located", [])
    if not_located:
        md.append("## 🔍 无法定位字段")
        md.append("| 字段名 | 原因 |")
        md.append("|--------|------|")
        for e in not_located:
            md.append(f"| {e.get('field_label','?')} | {e.get('reason','')} |")
        md.append("")

    marks = review.get("fields_marked_for_profile_save", [])
    if marks:
        md.append("## 💾 建议保存到资料库的字段")
        for e in marks:
            fk = e.get("field_key", "")
            val = mask_for_display(fk, e.get("value", ""))
            if len(val) > 30: val = val[:30] + "..."
            md.append(f"- {e.get('field_label','?')} ({fk}): {val}")
        md.append("> ⚠️ 以上字段仅记录，尚未写入 vault。如需保存请执行后续操作。")
        md.append("")

    if report.get("long_text_warning"):
        md.append("## ⚠️ 长文本复查提醒")
        md.append("以下字段为长文本，已写入最终文件，建议人工复查：")
        for e in report["long_text_warning"]:
            md.append(f"- {e.get('field_label','?')} → {e.get('location','?')}")
        md.append("")

    stats = report.get("_stats", {})
    md.append("---")
    md.append("## 📊 统计")
    md.append(f"| 指标 | 数值 |")
    md.append(f"|------|------|")
    md.append(f"| 成功写入 | {len(written)} |")
    md.append(f"| 跳过 | {len(skipped)} |")
    md.append(f"| 拒绝并清空 | {len(rejected)} |")
    md.append(f"| 无法定位 | {len(not_located)} |")
    md.append(f"| 修改原始文件 | 否 |")
    md.append(f"| 修改草稿文件 | 否 |")
    md.append(f"| 修改 vault | 否 |")
    md.append(f"| 联网 | 否 |")
    md.append(f"| 调用 API | 否 |")
    md.append("")

    mp = FINAL_REPORTS_DIR / f"final_report_{timestamp}.md"
    with open(mp, "w", encoding="utf-8") as f:
        f.write("\n".join(md))

    write_log(f"INFO: 最终报告已生成: {jp.name}, {mp.name}")
    return jp, mp


# ------------------------------------------------------------
# 主流程
# ------------------------------------------------------------
def main():
    print("\nSafeFill-FinalExport — 导出最终文件")
    print("不修改原始/草稿/vault | 不联网\n")
    write_log("========== 开始最终导出 ==========")

    # 1. 查找 review_result
    review_path = find_latest_review()
    if review_path is None:
        write_log("ERROR: 没有正式 review_result")
        print("\n[STOP] 没有找到正式网页检查结果。")
        print("请先在 SafeFill-FormReview 网页中点击【保存检查结果】。\n")
        sys.exit(1)

    try:
        with open(review_path, "r", encoding="utf-8") as f:
            review = json.load(f)
    except Exception as e:
        write_log(f"ERROR: 无法读取 {review_path.name}: {e}")
        print(f"\n错误：无法读取 {review_path.name}\n  {e}\n")
        sys.exit(1)

    # 有效性检查
    if not validate_review_freshness(review_path, review):
        sys.exit(1)
        sys.exit(1)

    write_log(f"INFO: 使用 review_result: {review_path.name}")

    # 2. 确认草稿文件 — 含路径边界校验
    draft_file_str = review.get("draft_file", "")
    try:
        draft_path = validate_draft_path(draft_file_str)
    except ValueError as e:
        print(f"\n错误：{e}\n")
        sys.exit(1)
    except (PermissionError, OSError):
        write_log(f"SECURITY: draft_file 路径越界，拒绝访问: {draft_file_str}")
        print(f"\n[STOP] 安全拒绝：draft_file 指向了 draft_outputs 目录之外。")
        print(f"  draft_file: {draft_file_str}")
        print(f"  允许范围: {DRAFT_DIR.resolve()}")
        print(f"  拒绝导出。\n")
        sys.exit(1)
    except FileNotFoundError as e:
        write_log(f"ERROR: {e}")
        print(f"\n错误：{e}")
        print(f"请确认草稿文件仍在 draft_outputs 目录中。\n")
        sys.exit(1)

    write_log(f"INFO: 草稿文件: {draft_path.name}")

    # 3. 确定输出格式
    ext = draft_path.suffix.lower()
    if ext not in (".docx", ".xlsx"):
        write_log(f"ERROR: 不支持的格式: {ext}")
        print(f"\n错误：不支持的格式 '{ext}'。仅支持 .docx 和 .xlsx。\n")
        sys.exit(1)

    # 导出前摘要
    show_export_summary(review_path, review, draft_path, ext)

    # 4. 确定最终文件名
    stem = draft_path.stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    final_path = make_final_path(draft_path, timestamp)
    final_name = final_path.name

    # 5. 复制草稿到最终文件
    try:
        shutil.copy2(str(draft_path), str(final_path))
        write_log(f"INFO: 复制草稿 -> {final_name}")
    except Exception as e:
        write_log(f"ERROR: 复制失败: {e}")
        print(f"错误：无法复制草稿文件：{e}\n")
        sys.exit(1)

    # 6. 导出
    print(f"  处理中: {final_name} ...")
    mode = review.get("mode", "legacy")
    if mode == "html_review":
        write_log("INFO: 检测到 html_review 模式，按网页编辑表格导出。")
        if ext == ".docx":
            export_report = export_html_review_docx(final_path, review)
        else:
            export_report = export_html_review_xlsx(final_path, review)
    else:
        if ext == ".docx":
            export_report = export_docx(final_path, final_path, review)
        else:
            export_report = export_xlsx(final_path, final_path, review)

    # 7. 生成报告
    original_file = review.get("source_file", "") or review.get("original_file", "")
    jp, mp = generate_final_reports(
        review_path, draft_path, final_path, original_file,
        export_report, review, timestamp
    )

    # 8. 输出摘要
    wc = len(export_report.get("written", []))
    sc = len(export_report.get("skipped", []))
    rc = len(export_report.get("rejected_cleared", []))
    nc = len(export_report.get("not_located", []))

    print(f"\n{'='*55}")
    print(f"  [OK] 第 4 阶段完成！")
    print(f"  最终文件: {final_path}")
    print(f"  写入: {wc} | 跳过: {sc} | 拒绝: {rc} | 未定位: {nc}")
    print(f"  报告: {mp.name}")
    print(f"")
    print(f"  [LOCK] 安全确认:")
    print(f"     - 原始文件: 未修改")
    print(f"     - 草稿文件: 未修改")
    print(f"     - vault:     未修改")
    print(f"     - 联网:      否")
    print(f"     - API:       否")
    print(f"{'='*55}")
    print(f"\n  本轮填表已完成。为减少隐私残留，可以使用 Cleaner 清理过程产物：")
    print(f"    python D:\\SafeFill\\app\\cleaner.py preview")
    print(f"  Cleaner 默认不清理 vault 和 final_outputs。")
    print(f"{'='*55}\n")

    write_log(f"========== 最终导出完成 | 写入{wc} 跳过{sc} 拒绝{rc} 未定位{nc} ==========")


if __name__ == "__main__":
    main()
