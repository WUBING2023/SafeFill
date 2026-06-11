# -*- coding: utf-8 -*-
"""
本地安全填表助手 - 自动填写草稿脚本
功能：读取 new_forms 中的新表格，使用 vault 资料库中已确认的个人资料，
     自动匹配并填写字段，生成草稿文件和填写报告。
安全约束：不联网、不修改原始文件、不修改资料库、报告中敏感字段脱敏。
"""

import os
import sys
import json
import re
import copy
import shutil
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

# ------------------------------------------------------------
# 项目根目录
# ------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent

NEW_FORMS_DIR = PROJECT_ROOT / "new_forms"
VAULT_PROFILE = PROJECT_ROOT / "vault" / "profile.json"
DRAFT_OUTPUT_DIR = PROJECT_ROOT / "draft_outputs"
REPORT_OUTPUT_DIR = PROJECT_ROOT / "filling_reports"
HTML_OUTPUT_DIR = PROJECT_ROOT / "review_html"
LOGS_DIR = PROJECT_ROOT / "logs"

# ------------------------------------------------------------
# 字段名到资料库 key 的映射（中文字段名 -> profile JSON key）
# ------------------------------------------------------------
FIELD_NAME_MAP = {
    "姓名": "name",
    "名字": "name",
    "申报人": "name",
    "申请人": "name",
    "负责人": "name",

    "性别": "gender",

    "出生日期": "birth_date",
    "出生年月": "birth_date",
    "生日": "birth_date",
    "出生年月日": "birth_date",

    "身份证号": "id_number",
    "身份证号码": "id_number",
    "证件号码": "id_number",
    "身份证": "id_number",
    "公民身份号码": "id_number",

    "手机号": "phone",
    "手机号码": "phone",
    "联系电话": "phone",
    "电话": "phone",
    "联系方式": "phone",
    "移动电话": "phone",

    "邮箱": "email",
    "电子邮件": "email",
    "电子邮箱": "email",
    "E-mail": "email",
    "Email": "email",
    "邮件": "email",

    "工作单位": "organization",
    "单位": "organization",
    "所在单位": "organization",
    "依托单位": "organization",
    "申报单位": "organization",
    "单位名称": "organization",

    "部门": "department",
    "院系": "department",
    "学院": "department",
    "所在部门": "department",
    "所在院系": "department",
    "系所": "department",

    "职称": "title",
    "专业技术职务": "title",
    "现任职称": "title",
    "职务": "title",

    "学历": "education",
    "最高学历": "education",
    "文化程度": "education",

    "学位": "degree",
    "最高学位": "degree",
    "学位名称": "degree",

    "专业": "major",
    "所学专业": "major",
    "攻读专业": "major",
    "学科专业": "major",
    "专业名称": "major",

    "专业方向": "research_area",
    "研究方向": "research_area",
    "研究领域": "research_area",
    "学科方向": "research_area",
    "从事专业": "research_area",

    "通讯地址": "address",
    "通信地址": "address",
    "联系地址": "address",
    "地址": "address",
    "家庭地址": "address",
    "住址": "address",
    "邮寄地址": "address",

    "个人简介": "biography",
    "个人简历": "biography",
    "简介": "biography",
    "简历": "biography",
    "个人陈述": "biography",
    "自我简介": "biography",

    "项目经历": "project_experience",
    "科研项目": "project_experience",
    "主持项目": "project_experience",
    "参与项目": "project_experience",
    "项目列表": "project_experience",
    "承担项目": "project_experience",
}

# 敏感字段
SENSITIVE_FIELDS = {"id_number", "phone", "address", "photo_path"}

# 长文本字段（需用户复核）
LONG_TEXT_FIELDS = {"biography", "project_experience"}
COMPLEX_KW = ["项目", "课题", "经历", "简介", "研究基础", "工作基础", "主要贡献", "承担任务", "申请理由", "代表成果", "科研成果", "获奖情况", "论文", "专利", "推广应用", "团队分工", "技术路线", "可行性", "创新点", "研究背景", "研究意义", "研究目标", "研究内容", "研究方案", "预期成果", "本人贡献", "社会服务", "学术成果", "发表论文"]

# 匹配阈值：字段名字符串相似度最小值
MIN_MATCH_LENGTH = 1


# ------------------------------------------------------------
# 工具函数
# ------------------------------------------------------------
def write_log(message: str):
    """写入本地日志，不记录敏感字段内容。"""
    LOGS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_file = LOGS_DIR / "fill_form_draft.log"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(f"[{timestamp}] {message}\n")
    print(f"[LOG] {message}")


def load_profile(profile_path: Path) -> dict:
    """加载用户资料库，提取所有已确认字段的值。"""
    if not profile_path.exists():
        return {}
    try:
        with open(profile_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        write_log(f"ERROR: 无法读取 {profile_path.name}: {e}")
        return {}

    # 提取字段值
    profile = {}
    for key, value in data.items():
        if key.startswith("_"):
            continue

        if isinstance(value, dict):
            actual = value.get("value", "")
        elif isinstance(value, list):
            actual = value if value else ""
        else:
            actual = value

        # 跳过空值
        if not actual or (isinstance(actual, str) and not actual.strip()):
            continue
        if isinstance(actual, list) and len(actual) == 0:
            continue

        profile[key] = actual

    # 加载自定义字段（保留中文原名作为 key）
    custom_fields = data.get("custom_fields", {})
    for cf_key, cf_val in custom_fields.items():
        if isinstance(cf_val, dict):
            cf_actual = cf_val.get("value", "")
        else:
            cf_actual = str(cf_val)
        if cf_actual and str(cf_actual).strip():
            profile[cf_key] = str(cf_actual)

    return profile


def match_field_name(cell_text: str) -> str | None:
    """
    尝试将单元格文本匹配到资料库字段 key。
    返回字段 key（如 "name"、"phone"），或 None。
    """
    text = cell_text.strip()
    if not text:
        return None

    if is_complex_field_label(text):
        return None

    # 精确匹配
    if text in FIELD_NAME_MAP:
        return FIELD_NAME_MAP[text]

    # 去掉冒号后匹配
    cleaned = re.sub(r'\s*[：:]\s*$', '', text).strip()
    if cleaned in FIELD_NAME_MAP:
        return FIELD_NAME_MAP[cleaned]

    # 包含匹配（字段名在文本中）
    for chinese_name, field_key in FIELD_NAME_MAP.items():
        if chinese_name in text and len(chinese_name) >= MIN_MATCH_LENGTH:
            # 避免过短匹配（如"位"匹配"单位"中的"位"）
            if len(chinese_name) >= 2 or chinese_name == text:
                return field_key

    return None


def is_complex_field_label(text: str) -> bool:
    """Return True for long semantic prompts that should not be filled by alias matching."""
    cleaned = re.sub(r'\s+', '', str(text).strip())
    if not cleaned:
        return False
    if cleaned in FIELD_NAME_MAP:
        return False
    if cleaned in {"常用个人简介", "常用项目经历"}:
        return False
    has_complex_keyword = any(kw in cleaned for kw in COMPLEX_KW)
    return has_complex_keyword and len(cleaned) >= 8


def is_empty_cell(value) -> bool:
    """判断单元格是否为空。"""
    if value is None:
        return True
    if isinstance(value, str):
        text = value.strip()
        if text == "":
            return True
        normalized = re.sub(r"\s+", "", text)
        placeholder_patterns = [
            r"^[_＿]+$",
            r"^[-－—–]+$",
            r"^[.。·•]+$",
            r"^[□☐]+$",
            r"^[/／]+$",
            r"^[（(]?\s*(无|暂无|待填|待填写|空|未填|不详|N/?A)\s*[）)]?$",
        ]
        if any(re.fullmatch(pattern, normalized, flags=re.IGNORECASE) for pattern in placeholder_patterns):
            return True
    return False


def _clean_label_text(text: str) -> str:
    """清理字段名单元格，用于匹配 custom_fields。"""
    cleaned = str(text).strip()
    cleaned = re.sub(r'\s*[：:]\s*$', '', cleaned).strip()
    cleaned = re.sub(r'[\s_＿\-—－\.·□]+$', '', cleaned).strip()
    return cleaned


def resolve_field_key(cell_text: str, profile: dict) -> str | None:
    """匹配标准字段或 vault custom_fields 中的中文自定义字段。"""
    field_key = match_field_name(cell_text)
    if field_key:
        return field_key

    cleaned = _clean_label_text(cell_text)
    if cleaned in profile:
        return cleaned
    if str(cell_text).strip() in profile:
        return str(cell_text).strip()
    return None


def get_profile_value(profile: dict, field_key: str | None, field_label: str):
    """按标准 key 或自定义中文字段名取值。"""
    if not field_key:
        return None, field_key
    fill_value = profile.get(field_key, None)
    if fill_value is not None:
        return fill_value, field_key
    cleaned = _clean_label_text(field_label)
    fill_value = profile.get(cleaned, None)
    if fill_value is not None:
        return fill_value, cleaned
    fill_value = profile.get(str(field_label).strip(), None)
    if fill_value is not None:
        return fill_value, str(field_label).strip()
    return None, field_key


def normalize_fill_value(value) -> str:
    """把资料库值转换成可写入表格的字符串。"""
    if isinstance(value, list):
        value = "\n".join(str(item) for item in value)
    return str(value).strip() if value is not None else ""


def parse_inline_blank_label(cell_text: str, profile: dict):
    """
    识别“姓名：____”或“姓名：”这种同单元格待填写格式。
    如果冒号后已有真实内容，不处理。
    """
    text = str(cell_text).strip()
    m = re.match(r'^(.+?)[：:]\s*(.*)$', text)
    if not m:
        return None
    label = m.group(1).strip()
    tail = m.group(2).strip()
    # 冒号后只允许空白、下划线、横线、方框等占位符。
    if re.sub(r'[\s_＿\-—－\.·□]+', '', tail):
        return None
    field_key = resolve_field_key(label, profile)
    if not field_key:
        return None
    return label, field_key


def inline_blank_has_placeholder(cell_text: str) -> bool:
    """判断“字段：____”冒号后是否有明确占位符。"""
    m = re.match(r'^(.+?)[：:]\s*(.*)$', str(cell_text).strip())
    if not m:
        return False
    tail = m.group(2).strip()
    return bool(tail)


def write_docx_cell(cell, value: str):
    """清空并写入 Word 单元格，尽量保留单元格本身格式。"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].add_run(value)
    else:
        cell.add_paragraph(value)


def make_filled_entry(field_label: str, field_key: str, value: str, **extra) -> dict:
    entry = {
        "field_label": field_label,
        "field_key": field_key,
        "value_length": len(value),
    }
    entry.update(extra)
    return entry


def format_entry_location(entry: dict) -> str:
    """兼容本地填写和 API 填写两种 report entry 结构。"""
    if not isinstance(entry, dict):
        return "-"
    if entry.get("location"):
        return str(entry.get("location"))
    row = entry.get("row")
    col = entry.get("col", entry.get("target_col"))
    if "sheet" in entry:
        loc = f"Sheet[{entry.get('sheet')}]"
    elif "table" in entry:
        loc = f"Table[{entry.get('table')}]"
    else:
        return "-"
    if row:
        loc += f" R{row}"
    if col:
        loc += f"C{col}"
    return loc


def can_write_xlsx_cell(ws, cell) -> tuple[bool, str]:
    """合并单元格保护：只允许写普通单元格或合并区域左上角。"""
    for merged_range in ws.merged_cells.ranges:
        if cell.coordinate in merged_range:
            if cell.column == merged_range.min_col and cell.row == merged_range.min_row:
                return True, ""
            return False, f"目标单元格在合并区域中（{merged_range}），不是左上角，避免破坏合并结构"
    return True, ""


# ------------------------------------------------------------
# 脱敏工具
# ------------------------------------------------------------
def mask_id_number(value: str) -> str:
    if not value:
        return value
    v = str(value).strip()
    if len(v) >= 15:
        return v[:3] + "*" * (len(v) - 7) + v[-4:]
    if len(v) >= 8:
        return v[:2] + "*" * (len(v) - 4) + v[-2:]
    return "***"


def mask_phone(value: str) -> str:
    if not value:
        return value
    v = str(value).strip().replace("-", "").replace(" ", "")
    if len(v) == 11 and v.isdigit():
        return v[:3] + "****" + v[-4:]
    if len(v) >= 8:
        return v[:3] + "****" + v[-2:]
    return "***"


def mask_address(value: str) -> str:
    if not value:
        return value
    v = str(value).strip()
    if len(v) <= 6:
        return v + "***"
    return v[:6] + "***"


def mask_sensitive_value(field_key: str, value) -> str:
    if not value:
        return ""
    s = str(value)
    if field_key == "id_number":
        return mask_id_number(s)
    elif field_key == "phone":
        return mask_phone(s)
    elif field_key == "address":
        return mask_address(s)
    return s


# ------------------------------------------------------------
# docx 处理
# ------------------------------------------------------------
def fill_docx(filepath: Path, profile: dict) -> dict:
    """
    处理 .docx 文件：
    1. 遍历所有表格，找到"字段名在左、值在右"的行。
    2. 若右侧为空，填入资料库中的对应值。
    3. 返回填写结果摘要。
    """
    import docx
    from docx.shared import Pt, RGBColor

    report = {
        "filled": [],        # 已填写字段
        "skipped": [],       # 未填写（无匹配资料）
        "review_needed": [], # 需用户复核
        "blocked": [],       # 已有内容未覆盖
        "unknown": [],       # 识别到字段名但不在资料库中
        "not_supported": [], # 不支持的内容
    }

    try:
        doc = docx.Document(str(filepath))
    except Exception as e:
        write_log(f"ERROR: 无法打开 {filepath.name}: {e}")
        report["not_supported"].append(f"无法打开文件: {e}")
        return report

    tables_processed = 0
    cells_checked = 0
    cells_filled = 0

    for table_idx, table in enumerate(doc.tables):
        tables_processed += 1
        rows = table.rows

        for row_idx, row in enumerate(rows):
            cells = row.cells
            cell_count = len(cells)

            # 快照：处理前保存所有单元格原始文本
            cell_texts_before = [c.text.strip() for c in cells]

            # 按 pair 处理：col 和 col+1 为一组
            col_idx = 0
            inline_blank_count = sum(1 for t in cell_texts_before if parse_inline_blank_label(t, profile))
            while col_idx < cell_count:
                # 跳过空白字段名单元格
                if not cell_texts_before[col_idx]:
                    col_idx += 1
                    continue

                cell_text = cell_texts_before[col_idx]
                cells_checked += 1

                # 新模式 A：同一单元格中“姓名：____”或同一行多个“字段：”
                inline = parse_inline_blank_label(cell_text, profile)
                if inline and (inline_blank_has_placeholder(cell_text) or inline_blank_count > 1 or col_idx + 1 >= cell_count):
                    label, inline_key = inline
                    fill_value, resolved_key = get_profile_value(profile, inline_key, label)
                    fill_value_str = normalize_fill_value(fill_value)
                    if fill_value_str:
                        try:
                            write_docx_cell(cells[col_idx], f"{label}：{fill_value_str}")
                            cells_filled += 1
                            entry = make_filled_entry(
                                label, resolved_key, fill_value_str,
                                table=table_idx + 1,
                                row=row_idx + 1,
                                target_col=col_idx + 1,
                                fill_mode="inline_colon",
                            )
                            if resolved_key in LONG_TEXT_FIELDS:
                                entry["review_needed"] = True
                                report["review_needed"].append(entry)
                            else:
                                report["filled"].append(entry)
                            write_log(f"INFO: 已填写 {filepath.name} -> {label} (同单元格冒号模式)")
                        except Exception as e:
                            write_log(f"WARN: 写入同单元格失败 {filepath.name} -> {label}: {e}")
                            report["not_supported"].append(f"写入失败 [{label}]: {e}")
                    col_idx += 1
                    continue

                # 尝试匹配字段名（标准 + 自定义）
                field_key = resolve_field_key(cell_text, profile)
                if field_key is None:
                    if is_complex_field_label(cell_text):
                        report["unknown"].append({
                            "table": table_idx + 1,
                            "row": row_idx + 1,
                            "field_label": cell_text,
                            "field_key": "complex_unknown",
                        })
                        col_idx += 2
                        continue
                    if cell_text in profile:
                        field_key = cell_text
                    else:
                        col_idx += 1
                        continue

                # 检查右侧单元格
                target_col = col_idx + 1
                if target_col >= cell_count:
                    # 新模式 A：同一单元格中“姓名：____”或“姓名：”
                    inline = parse_inline_blank_label(cell_text, profile)
                    if inline:
                        label, inline_key = inline
                        fill_value, resolved_key = get_profile_value(profile, inline_key, label)
                        fill_value_str = normalize_fill_value(fill_value)
                        if fill_value_str:
                            try:
                                write_docx_cell(cells[col_idx], f"{label}：{fill_value_str}")
                                cells_filled += 1
                                entry = make_filled_entry(
                                    label, resolved_key, fill_value_str,
                                    table=table_idx + 1,
                                    row=row_idx + 1,
                                    target_col=col_idx + 1,
                                    fill_mode="inline_colon",
                                )
                                if resolved_key in LONG_TEXT_FIELDS:
                                    entry["review_needed"] = True
                                    report["review_needed"].append(entry)
                                else:
                                    report["filled"].append(entry)
                                write_log(f"INFO: 已填写 {filepath.name} -> {label} (同单元格冒号模式)")
                            except Exception as e:
                                write_log(f"WARN: 写入同单元格失败 {filepath.name} -> {label}: {e}")
                                report["not_supported"].append(f"写入失败 [{label}]: {e}")
                    col_idx += 2
                    continue

                target_text = cell_texts_before[target_col] if target_col < len(cell_texts_before) else ""

                # 如果右侧已有内容，不覆盖
                if not is_empty_cell(target_text):
                    report["blocked"].append({
                        "table": table_idx + 1,
                        "row": row_idx + 1,
                        "field_label": cell_text,
                        "field_key": field_key,
                        "existing_value": target_text[:50],
                        "reason": "右侧单元格已有内容，未覆盖",
                    })
                    col_idx += 2
                    continue

                # 查资料库
                fill_value, field_key = get_profile_value(profile, field_key, cell_text)
                if fill_value is None:
                    report["unknown"].append({
                        "table": table_idx + 1,
                        "row": row_idx + 1,
                        "field_label": cell_text,
                        "field_key": field_key,
                    })
                    col_idx += 2
                    continue

                fill_value_str = normalize_fill_value(fill_value)
                if not fill_value_str:
                    report["skipped"].append({
                        "table": table_idx + 1,
                        "row": row_idx + 1,
                        "field_label": cell_text,
                        "field_key": field_key,
                        "reason": "资料库中值为空",
                    })
                    col_idx += 2
                    continue

                # 写入
                try:
                    target_cell = cells[target_col]
                    write_docx_cell(target_cell, fill_value_str)

                    cells_filled += 1
                    entry = make_filled_entry(
                        cell_text, field_key, fill_value_str,
                        table=table_idx + 1,
                        row=row_idx + 1,
                        label_col=col_idx + 1,
                        target_col=target_col + 1,
                        fill_mode="right_cell",
                    )
                    if field_key in LONG_TEXT_FIELDS:
                        entry["review_needed"] = True
                        report["review_needed"].append(entry)
                    else:
                        report["filled"].append(entry)

                    write_log(f"INFO: 已填写 {filepath.name} -> {cell_text} (字段: {field_key})")

                except Exception as e:
                    write_log(f"WARN: 写入单元格失败 {filepath.name} -> {cell_text}: {e}")
                    report["not_supported"].append(f"写入失败 [{cell_text}]: {e}")

                col_idx += 2  # 跳到下一 pair

        # 新模式 B：字段在上一行，值在下一行同列
        for upper_row_idx in range(len(rows) - 1):
            upper_cells = rows[upper_row_idx].cells
            lower_cells = rows[upper_row_idx + 1].cells
            max_cols = min(len(upper_cells), len(lower_cells))
            for col_idx in range(max_cols):
                label_text = upper_cells[col_idx].text.strip()
                if not label_text:
                    continue
                if "：" in label_text or ":" in label_text:
                    continue
                field_key = resolve_field_key(label_text, profile)
                if not field_key:
                    continue
                lower_text = lower_cells[col_idx].text.strip()
                if not is_empty_cell(lower_text):
                    continue
                fill_value, field_key = get_profile_value(profile, field_key, label_text)
                fill_value_str = normalize_fill_value(fill_value)
                if not fill_value_str:
                    continue
                try:
                    write_docx_cell(lower_cells[col_idx], fill_value_str)
                    cells_filled += 1
                    entry = make_filled_entry(
                        label_text, field_key, fill_value_str,
                        table=table_idx + 1,
                        row=upper_row_idx + 2,
                        label_row=upper_row_idx + 1,
                        target_col=col_idx + 1,
                        fill_mode="below_cell",
                    )
                    if field_key in LONG_TEXT_FIELDS:
                        entry["review_needed"] = True
                        report["review_needed"].append(entry)
                    else:
                        report["filled"].append(entry)
                    write_log(f"INFO: 已填写 {filepath.name} -> {label_text} (字段在上值在下)")
                except Exception as e:
                    write_log(f"WARN: 写入下方单元格失败 {filepath.name} -> {label_text}: {e}")
                    report["not_supported"].append(f"写入失败 [{label_text}]: {e}")

    report["_stats"] = {
        "tables_processed": tables_processed,
        "cells_checked": cells_checked,
        "cells_filled": cells_filled,
    }
    # 保存修改到磁盘
    doc.save(str(filepath))
    return report


# ------------------------------------------------------------
# xlsx 处理
# ------------------------------------------------------------
def fill_xlsx(filepath: Path, profile: dict) -> dict:
    """
    处理 .xlsx 文件：
    1. 遍历所有工作表，找到"字段名在左、值在右"的单元格对。
    2. 若右侧为空，填入资料库中的对应值。
    3. 尽量保留格式（通过 openpyxl 复制样式）。
    4. 返回填写结果摘要。
    """
    import openpyxl

    report = {
        "filled": [],
        "skipped": [],
        "review_needed": [],
        "blocked": [],
        "unknown": [],
        "not_supported": [],
    }

    try:
        wb = openpyxl.load_workbook(str(filepath))
    except Exception as e:
        write_log(f"ERROR: 无法打开 {filepath.name}: {e}")
        report["not_supported"].append(f"无法打开文件: {e}")
        return report

    sheets_processed = 0
    cells_checked = 0
    cells_filled = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheets_processed += 1

        # 扫描所有行
        for row in ws.iter_rows():
            row_cells = [cell for cell in row]
            cell_count = len(row_cells)
            # 快照
            cell_texts_before = [str(c.value).strip() if c.value is not None else "" for c in row_cells]

            col_idx = 0
            inline_blank_count = sum(1 for t in cell_texts_before if parse_inline_blank_label(t, profile))
            while col_idx < cell_count:
                if not cell_texts_before[col_idx]:
                    col_idx += 1
                    continue
                cell_text = cell_texts_before[col_idx]
                cells_checked += 1

                # 新模式 A：同一单元格中“姓名：____”或同一行多个“字段：”
                inline = parse_inline_blank_label(cell_text, profile)
                if inline and (inline_blank_has_placeholder(cell_text) or inline_blank_count > 1 or col_idx + 1 >= cell_count):
                    label, inline_key = inline
                    fill_value, resolved_key = get_profile_value(profile, inline_key, label)
                    fill_value_str = normalize_fill_value(fill_value)
                    if fill_value_str:
                        ok, reason = can_write_xlsx_cell(ws, row_cells[col_idx])
                        if ok:
                            try:
                                row_cells[col_idx].value = f"{label}：{fill_value_str}"
                                cells_filled += 1
                                entry = make_filled_entry(
                                    label, resolved_key, fill_value_str,
                                    sheet=sheet_name,
                                    row=row_cells[col_idx].row,
                                    col=row_cells[col_idx].column,
                                    fill_mode="inline_colon",
                                )
                                if resolved_key in LONG_TEXT_FIELDS:
                                    entry["review_needed"] = True
                                    report["review_needed"].append(entry)
                                else:
                                    report["filled"].append(entry)
                                write_log(f"INFO: 已填写 {filepath.name} -> {label} (同单元格冒号模式)")
                            except Exception as e:
                                write_log(f"WARN: 写入同单元格失败 {filepath.name} -> {label}: {e}")
                                report["not_supported"].append(f"写入失败 [{label}]: {e}")
                        else:
                            report["not_supported"].append({
                                "sheet": sheet_name,
                                "row": row_cells[col_idx].row,
                                "field_label": label,
                                "reason": reason,
                            })
                    col_idx += 1
                    continue

                # 尝试匹配字段名（标准 + 自定义）
                field_key = resolve_field_key(cell_text, profile)
                if field_key is None:
                    if is_complex_field_label(cell_text):
                        report["unknown"].append({
                            "sheet": sheet_name,
                            "row": row_cells[col_idx].row,
                            "field_label": cell_text,
                            "field_key": "complex_unknown",
                        })
                        col_idx += 2
                        continue
                    if cell_text in profile:
                        field_key = cell_text
                    else:
                        col_idx += 1
                        continue

                # 检查右侧单元格
                target_col = col_idx + 1
                if target_col >= cell_count:
                    # 新模式 A：同一单元格中“姓名：____”或“姓名：”
                    inline = parse_inline_blank_label(cell_text, profile)
                    if inline:
                        label, inline_key = inline
                        fill_value, resolved_key = get_profile_value(profile, inline_key, label)
                        fill_value_str = normalize_fill_value(fill_value)
                        if fill_value_str:
                            ok, reason = can_write_xlsx_cell(ws, row_cells[col_idx])
                            if ok:
                                try:
                                    row_cells[col_idx].value = f"{label}：{fill_value_str}"
                                    cells_filled += 1
                                    entry = make_filled_entry(
                                        label, resolved_key, fill_value_str,
                                        sheet=sheet_name,
                                        row=row_cells[col_idx].row,
                                        col=row_cells[col_idx].column,
                                        fill_mode="inline_colon",
                                    )
                                    if resolved_key in LONG_TEXT_FIELDS:
                                        entry["review_needed"] = True
                                        report["review_needed"].append(entry)
                                    else:
                                        report["filled"].append(entry)
                                    write_log(f"INFO: 已填写 {filepath.name} -> {label} (同单元格冒号模式)")
                                except Exception as e:
                                    write_log(f"WARN: 写入同单元格失败 {filepath.name} -> {label}: {e}")
                                    report["not_supported"].append(f"写入失败 [{label}]: {e}")
                            else:
                                report["not_supported"].append({
                                    "sheet": sheet_name,
                                    "row": row_cells[col_idx].row,
                                    "field_label": label,
                                    "reason": reason,
                                })
                    col_idx += 2
                    continue

                target_cell = row_cells[target_col]
                target_text = cell_texts_before[target_col] if target_col < len(cell_texts_before) else ""

                # 如果右侧已有内容，不覆盖
                if not is_empty_cell(target_text):
                    existing = str(target_text).strip()[:50]
                    report["blocked"].append({
                        "sheet": sheet_name,
                        "row": target_cell.row,
                        "col": target_cell.column,
                        "field_label": cell_text,
                        "field_key": field_key,
                        "existing_value": existing,
                        "reason": "右侧单元格已有内容，未覆盖",
                    })
                    col_idx += 2
                    continue

                # 检查是否为合并单元格（不破坏合并结构）
                ok, merge_reason = can_write_xlsx_cell(ws, target_cell)
                if not ok:
                    report["not_supported"].append({
                        "sheet": sheet_name,
                        "row": target_cell.row,
                        "field_label": cell_text,
                        "reason": merge_reason,
                    })
                    col_idx += 2
                    continue

                # 查资料库（标准字段 + 自定义字段）
                fill_value, field_key = get_profile_value(profile, field_key, cell_text)
                if fill_value is None:
                    report["unknown"].append({
                        "sheet": sheet_name,
                        "row": target_cell.row,
                        "field_label": cell_text,
                        "field_key": field_key,
                    })
                    col_idx += 2
                    continue

                fill_value_str = normalize_fill_value(fill_value)
                if not fill_value_str:
                    report["skipped"].append({
                        "sheet": sheet_name,
                        "row": target_cell.row,
                        "field_label": cell_text,
                        "field_key": field_key,
                        "reason": "资料库中值为空",
                    })
                    col_idx += 2
                    continue

                # 写入目标单元格
                try:
                    target_cell.value = fill_value_str
                    cells_filled += 1

                    entry = make_filled_entry(
                        cell_text, field_key, fill_value_str,
                        sheet=sheet_name,
                        row=target_cell.row,
                        col=target_cell.column,
                        fill_mode="right_cell",
                    )

                    if field_key in LONG_TEXT_FIELDS:
                        entry["review_needed"] = True
                        report["review_needed"].append(entry)
                    else:
                        report["filled"].append(entry)

                    write_log(f"INFO: 已填写 {filepath.name} -> {cell_text} (字段: {field_key})")

                except Exception as e:
                    write_log(f"WARN: 写入单元格失败 {filepath.name} -> {cell_text}: {e}")
                    report["not_supported"].append(f"写入失败 [{cell_text}]: {e}")

                col_idx += 2  # 下一 pair

        # 新模式 B：字段在上一行，值在下一行同列
        max_row_for_vertical = ws.max_row
        max_col_for_vertical = ws.max_column
        for upper_row_num in range(1, max_row_for_vertical + 1):
            lower_row_num = upper_row_num + 1
            for col_num in range(1, max_col_for_vertical + 1):
                upper_cell = ws.cell(row=upper_row_num, column=col_num)
                lower_cell = ws.cell(row=lower_row_num, column=col_num)
                label_text = str(upper_cell.value).strip() if upper_cell.value is not None else ""
                if not label_text:
                    continue
                if "：" in label_text or ":" in label_text:
                    continue
                field_key = resolve_field_key(label_text, profile)
                if not field_key:
                    continue
                lower_text = str(lower_cell.value).strip() if lower_cell.value is not None else ""
                if not is_empty_cell(lower_text):
                    continue
                ok, reason = can_write_xlsx_cell(ws, lower_cell)
                if not ok:
                    report["not_supported"].append({
                        "sheet": sheet_name,
                        "row": lower_cell.row,
                        "field_label": label_text,
                        "reason": reason,
                    })
                    continue
                fill_value, field_key = get_profile_value(profile, field_key, label_text)
                fill_value_str = normalize_fill_value(fill_value)
                if not fill_value_str:
                    continue
                try:
                    lower_cell.value = fill_value_str
                    cells_filled += 1
                    entry = make_filled_entry(
                        label_text, field_key, fill_value_str,
                        sheet=sheet_name,
                        row=lower_cell.row,
                        col=lower_cell.column,
                        label_row=upper_cell.row,
                        fill_mode="below_cell",
                    )
                    if field_key in LONG_TEXT_FIELDS:
                        entry["review_needed"] = True
                        report["review_needed"].append(entry)
                    else:
                        report["filled"].append(entry)
                    write_log(f"INFO: 已填写 {filepath.name} -> {label_text} (字段在上值在下)")
                except Exception as e:
                    write_log(f"WARN: 写入下方单元格失败 {filepath.name} -> {label_text}: {e}")
                    report["not_supported"].append(f"写入失败 [{label_text}]: {e}")

    report["_stats"] = {
        "sheets_processed": sheets_processed,
        "cells_checked": cells_checked,
        "cells_filled": cells_filled,
    }
    # 保存修改到磁盘
    wb.save(str(filepath))
    return report


# ------------------------------------------------------------
# 报告生成
# ------------------------------------------------------------
def format_entry_location(entry: dict) -> str:
    """兼容本地填写和 API 填写两种 report entry 结构。"""
    if not isinstance(entry, dict):
        return "-"
    if entry.get("location"):
        return str(entry.get("location"))
    row = entry.get("row")
    col = entry.get("col", entry.get("target_col"))
    if "sheet" in entry:
        loc = f"Sheet[{entry.get('sheet')}]"
    elif "table" in entry:
        loc = f"Table[{entry.get('table')}]"
    else:
        return "-"
    if row:
        loc += f" R{row}"
    if col:
        loc += f"C{col}"
    return loc


def generate_reports(filepath: Path, draft_path: Path, profile_path: Path,
                     report: dict, profile: dict, timestamp: str):
    """生成 JSON 和 Markdown 填写报告。文件名包含源文件名防止覆盖。"""

    # 安全文件名：取源文件 stem，替换特殊字符
    safe_stem = re.sub(r'[\\/:*?"<>|]', '_', filepath.stem)[:40]

    # ----- JSON 报告 -----
    base_name = f"filling_report_{safe_stem}_{timestamp}"
    json_filename = f"{base_name}.json"
    json_path = REPORT_OUTPUT_DIR / json_filename
    # 如果同名文件存在（极端情况），追加递增编号
    counter = 1
    while json_path.exists():
        counter += 1
        json_filename = f"{base_name}_{counter:03d}.json"
        json_path = REPORT_OUTPUT_DIR / json_filename

    json_data = {
        "_说明": "填写报告 - 仅供本地使用，不得联网发送。",
        "_生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "原始文件": str(filepath),
        "草稿文件": str(draft_path),
        "资料档案": str(profile_path),
        "是否联网": False,
        "是否调用API": False,
        "已填写字段数": len(report.get("filled", [])),
        "未填写字段数": len(report.get("skipped", [])) + len(report.get("unknown", [])),
        "需复核字段数": len(report.get("review_needed", [])),
        "已阻塞字段数": len(report.get("blocked", [])),
        "已填写字段": report.get("filled", []),
        "未填写字段": report.get("skipped", []) + report.get("unknown", []),
        "complex_unknown_fields": [
            {**f, "api_preview_allowed": True, "sensitive": False}
            for f in (report.get("unknown", []) + report.get("skipped", []))
            if isinstance(f, dict) and any(kw in str(f.get("field_label", "")) for kw in COMPLEX_KW)
        ],
        "需用户复核字段": report.get("review_needed", []),
        "已有内容未覆盖": report.get("blocked", []),
        "不支持或跳过": report.get("not_supported", []),
        "是否发现敏感字段": any(
            f.get("field_key") in SENSITIVE_FIELDS
            for f in report.get("filled", []) + report.get("review_needed", [])
        ),
        "_stats": report.get("_stats", {}),
    }

    REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)

    # ----- Markdown 报告（脱敏） -----
    md_filename = f"{base_name}.md"
    md_path = REPORT_OUTPUT_DIR / md_filename
    counter = 1
    while md_path.exists():
        counter += 1
        md_filename = f"{base_name}_{counter:03d}.md"
        md_path = REPORT_OUTPUT_DIR / md_filename

    md = []
    md.append("# 自动填写报告")
    md.append("")
    md.append(f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    md.append(f"> 原始文件：{filepath.name}")
    md.append(f"> 草稿文件：{draft_path.name}")
    md.append(f"> 资料档案：{profile_path.name}")
    md.append(f"> 联网：否 | API：否")
    md.append("")
    md.append("---")
    md.append("")

    # 已填写字段
    filled = report.get("filled", [])
    review_needed = report.get("review_needed", [])
    all_filled = filled + review_needed

    if all_filled:
        md.append("## ✅ 已填写字段")
        md.append("")
        md.append("| 字段名称 | 资料库字段 | 填写值（脱敏） | 位置 | 备注 |")
        md.append("|----------|-----------|---------------|------|------|")
        for entry in all_filled:
            label = entry.get("field_label", "?")
            fkey = entry.get("field_key", "?")
            raw_value = profile.get(fkey, "")
            # 脱敏
            display_val = mask_sensitive_value(fkey, raw_value) if fkey in SENSITIVE_FIELDS else str(raw_value)
            if len(display_val) > 40:
                display_val = display_val[:40] + "..."
            display_val = display_val.replace("|", "/")

            # 位置（兼容本地和 API 结构）
            loc = format_entry_location(entry)

            note = "⚠️ 需用户复核" if entry in review_needed else ""
            md.append(f"| {label} | {fkey} | {display_val} | {loc} | {note} |")
        md.append("")

    # 未填写字段
    skipped = report.get("skipped", [])
    unknown = report.get("unknown", [])
    all_missed = skipped + unknown

    if all_missed:
        md.append("## ❌ 未填写字段")
        md.append("")
        md.append("| 字段名称 | 资料库字段 | 原因 | 位置 |")
        md.append("|----------|-----------|------|------|")
        for entry in all_missed:
            label = entry.get("field_label", "?")
            fkey = entry.get("field_key", "?")
            reason = entry.get("reason", "资料库中无此信息")
            loc = format_entry_location(entry)
            md.append(f"| {label} | {fkey} | {reason} | {loc} |")
        md.append("")

    # 已有内容未覆盖
    blocked = report.get("blocked", [])
    if blocked:
        md.append("## 🔒 已有内容未覆盖")
        md.append("")
        md.append("| 字段名称 | 已有内容 | 位置 |")
        md.append("|----------|---------|------|")
        for entry in blocked:
            label = entry.get("field_label", "?")
            existing = entry.get("existing_value", "")
            existing = mask_sensitive_value(entry.get("field_key", ""), existing) if entry.get("field_key") in SENSITIVE_FIELDS else existing
            if len(existing) > 30:
                existing = existing[:30] + "..."
            loc = format_entry_location(entry)
            md.append(f"| {label} | {existing} | {loc} |")
        md.append("")

    # 不支持或跳过
    not_supported = report.get("not_supported", [])
    if not_supported:
        md.append("## ⚠️ 不支持或跳过")
        md.append("")
        for item in not_supported:
            if isinstance(item, str):
                md.append(f"- {item}")
            else:
                md.append(f"- {item.get('field_label', '?')}: {item.get('reason', str(item))}")
        md.append("")

    # 统计信息
    stats = report.get("_stats", {})
    md.append("---")
    md.append("")
    md.append("## 📊 统计")
    md.append("")
    md.append(f"| 指标 | 数值 |")
    md.append(f"|------|------|")
    md.append(f"| 已填写字段 | {len(filled)} |")
    md.append(f"| 需复核字段 | {len(review_needed)} |")
    md.append(f"| 未填写字段 | {len(all_missed)} |")
    md.append(f"| 已有内容未覆盖 | {len(blocked)} |")
    md.append(f"| 是否联网 | 否 |")
    md.append(f"| 是否调用 API | 否 |")
    md.append(f"| 是否发现敏感字段 | {'是（已在报告中脱敏）' if json_data['是否发现敏感字段'] else '否'} |")
    if "cells_filled" in stats:
        md.append(f"| 检查单元格数 | {stats.get('cells_checked', '?')} |")
        md.append(f"| 实际填写单元格数 | {stats.get('cells_filled', '?')} |")
    md.append("")
    md.append("---")
    md.append("")
    md.append("## ⚠️ 重要提示")
    md.append("")
    md.append("- **本文件是草稿，不是最终版。** 请在第 3 阶段（本地网页检查）中核对。")
    md.append("- 敏感信息（身份证号、手机号、地址）在本报告中已脱敏显示。")
    md.append("- 标记'需用户复核'的字段（项目经历、个人简介等）请仔细核对。")
    md.append("- 未填写的字段需要在下一阶段手动补充。")
    md.append("- 本报告未联网发送，仅保存在本地。")
    md.append("")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md))

    write_log(f"INFO: 已生成填写报告: {md_filename}, {json_filename}")

    return json_path, md_path


# ------------------------------------------------------------
# HTML 预览生成
# ------------------------------------------------------------
def generate_html_preview(filepath: Path, draft_path: Path, profile: dict, report: dict, timestamp: str) -> Path:
    """生成可编辑 HTML 表格预览。"""
    import html as html_mod
    HTML_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    stem = re.sub(r'[\\/:*?"<>|]', '_', filepath.stem)[:40]
    html_name = f"{stem}_{timestamp}.html"
    html_path = HTML_OUTPUT_DIR / html_name

    filled = report.get("filled", [])
    review_needed = report.get("review_needed", [])
    filled_targets = set()
    for f in filled + review_needed:
        if "table" in f and "row" in f:
            col = f.get("target_col", f.get("col"))
            if col:
                filled_targets.add((f.get("table"), f.get("row"), col))
        if "sheet" in f and "row" in f:
            col = f.get("col", f.get("target_col"))
            if col:
                filled_targets.add((f.get("sheet"), f.get("row"), col))

    ext = filepath.suffix.lower()
    tables_html = ""
    if ext == ".docx":
        import docx as dx
        try:
            doc = dx.Document(str(draft_path))
            for ti, table in enumerate(doc.tables):
                rows_html = ""
                for ri, row in enumerate(table.rows):
                    cells_html = ""
                    for ci, cell in enumerate(row.cells):
                        ct = cell.text.strip()
                        is_filled_target = (ti + 1, ri + 1, ci + 1) in filled_targets
                        bg = "#e8f5e9" if is_filled_target else ("#fff9c4" if not ct else "#ffffff")
                        display = html_mod.escape(ct)
                        cells_html += f'<td style="padding:4px 8px;border:1px solid #ccc;background:{bg}" contenteditable="true">{display}</td>'
                    rows_html += f"<tr>{cells_html}</tr>"
                tables_html += f'<h3>Table {ti+1}</h3><table style="border-collapse:collapse;width:100%">{rows_html}</table><br>'
        except Exception:
            tables_html = "<p>无法生成表格预览</p>"

    elif ext == ".xlsx":
        import openpyxl
        try:
            wb = openpyxl.load_workbook(str(draft_path), data_only=True)
            for sn in wb.sheetnames:
                ws = wb[sn]
                tables_html += f"<h3>{html_mod.escape(sn)}</h3><table style='border-collapse:collapse;width:100%'>"
                for row in ws.iter_rows():
                    cells_html = ""
                    for cell in row:
                        ct = str(cell.value).strip() if cell.value is not None else ""
                        is_filled_target = (sn, cell.row, cell.column) in filled_targets
                        bg = "#e8f5e9" if is_filled_target else ("#fff9c4" if not ct else "#ffffff")
                        cells_html += f'<td style="padding:4px 8px;border:1px solid #ccc;background:{bg}" contenteditable="true">{html_mod.escape(ct)}</td>'
                    tables_html += f"<tr>{cells_html}</tr>"
                tables_html += "</table><br>"
            wb.close()
        except Exception:
            tables_html = "<p>无法生成表格预览</p>"

    filled_count = len(filled)
    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="UTF-8"><title>SafeFill 可编辑预览</title>
<style>
body{{font-family:"Microsoft YaHei",sans-serif;max-width:900px;margin:20px auto;padding:0 20px}}
h2{{color:#1976d2}} .bar{{background:#1976d2;color:#fff;padding:8px 16px;border-radius:6px;margin-bottom:16px;display:flex;justify-content:space-between;align-items:center}}
td:focus{{outline:2px solid #1976d2}} .hint{{color:#666;font-size:13px;margin-top:8px}}
button{{padding:8px 16px;border:1px solid #ccc;border-radius:4px;cursor:pointer;background:#1976d2;color:#fff;font-size:14px}}
button:hover{{opacity:0.9}}
</style></head>
<body>
<div class="bar"><span>SafeFill 可编辑预览</span><span>{html_mod.escape(filepath.name)} | 已填 {filled_count} 字段</span></div>
{tables_html}
<div class="hint">浅绿=已自动填写的值 | 浅黄=空白区(可编辑) | 修改后点击下方按钮保存</div>
<button onclick="saveEdits()">保存修改为 JSON</button>
<p class="hint">这是本地 HTML 预览，修改后需保存 JSON，后续由 FinalExport 导出正式 Word/Excel。不联网。</p>
<script>
function saveEdits(){{var tables=document.querySelectorAll('table');var data={{tables:[]}};
tables.forEach(function(t,i){{var rows=[];t.querySelectorAll('tr').forEach(function(r){{var cells=[];r.querySelectorAll('td').forEach(function(c){{cells.push(c.textContent.trim())}});rows.push(cells)}});data.tables.push({{table_index:i+1,rows:rows}})}});
data.source_file="{html_mod.escape(str(filepath))}";data.draft_file="{html_mod.escape(str(draft_path))}";
var blob=new Blob([JSON.stringify(data,null,2)],{{type:'application/json'}});
var a=document.createElement('a');a.href=URL.createObjectURL(blob);a.download='review_edit_{timestamp}.json';a.click()
}}
</script></body></html>"""

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)

    # 生成配套 JSON
    json_name = f"{stem}_{timestamp}.json"
    json_path = HTML_OUTPUT_DIR / json_name
    json_data = {
        "source_file": str(filepath),
        "draft_file": str(draft_path),
        "html_file": str(html_path),
        "filled_count": filled_count,
        "filled_fields": [f.get("field_label", "") for f in filled],
    }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)

    return html_path


def write_latest_review_html(entries: list):
    """生成 latest_review_html.json 指针文件。"""
    HTML_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    latest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "items": entries,
    }
    with open(HTML_OUTPUT_DIR / "latest_review_html.json", "w", encoding="utf-8") as f:
        json.dump(latest, f, ensure_ascii=False, indent=2)


# ------------------------------------------------------------
# API-first 表单结构提取
# ------------------------------------------------------------
# ------------------------------------------------------------
# Word 文本框结构识别（供 API-first form_structure 使用）
# ------------------------------------------------------------
TEXTBOX_LABEL_KEYWORDS = [
    "个人简介", "项目经历", "本人贡献", "研究方向", "备注",
    "通讯地址", "联系方式", "导师", "学号", "民族", "政治面貌",
    "教育背景", "工作经历", "获奖情况", "资格证书", "爱好",
    "擅长方向", "自我评价", "职业技能",
]

TEXTBOX_XML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}


def _guess_textbox_label(text: str) -> str:
    """Guess what field label a textbox represents based on keyword matching."""
    for kw in TEXTBOX_LABEL_KEYWORDS:
        if kw in text:
            return kw
    return ""


def extract_docx_textbox_structure(filepath: Path) -> list:
    """
    Extract Word textbox/shape structure from a blank form .docx.
    Returns list of dicts with type, textbox_index, text, is_empty, label_guess, chars.
    """
    result = []
    try:
        with zipfile.ZipFile(str(filepath), "r") as zf:
            if "word/document.xml" not in zf.namelist():
                return result
            xml_bytes = zf.read("word/document.xml")
    except Exception:
        return result

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return result

    texts = []

    def _tb_paragraphs(txb_el):
        lines = []
        for p_el in txb_el.iter(f'{{{TEXTBOX_XML_NS["w"]}}}p'):
            t_nodes = p_el.iter(f'{{{TEXTBOX_XML_NS["w"]}}}t')
            line = "".join(t.text or "" for t in t_nodes).strip()
            if line:
                lines.append(line)
        return "\n".join(lines) if lines else ""

    # w:txbxContent
    for el in root.iter(f'{{{TEXTBOX_XML_NS["w"]}}}txbxContent'):
        text = _tb_paragraphs(el)
        if text:
            texts.append(text)

    # v:textbox
    for el in root.iter(f'{{{TEXTBOX_XML_NS["v"]}}}textbox'):
        text = _tb_paragraphs(el)
        if text:
            texts.append(text)

    # wps:txbx
    for el in root.iter(f'{{{TEXTBOX_XML_NS["wps"]}}}txbx'):
        text = _tb_paragraphs(el)
        if text:
            texts.append(text)

    # Dedup: keep first occurrence of each non-empty text
    seen = set()
    deduped = []
    for text in texts:
        normalized = text.strip()
        if normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(text)

    for i, text in enumerate(deduped):
        result.append({
            "type": "docx_textbox",
            "textbox_index": i + 1,
            "text": text,
            "is_empty": len(text) == 0,
            "label_guess": _guess_textbox_label(text),
            "chars": len(text),
        })

    return result


def extract_docx_form_structure(filepath: Path) -> dict:
    """提取 docx 空白表结构供 API 使用。"""
    import docx
    try:
        doc = docx.Document(str(filepath))
    except Exception as e:
        return {"file_name": filepath.name, "file_type": "docx", "error": str(e)}
    paras = [{"index": i+1, "text": p.text.strip()} for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = [{"col_index": ci+1, "text": cell.text.strip()} for ci, cell in enumerate(row.cells)]
            rows.append({"row_index": ri+1, "cells": cells})
        tables.append({"table_index": ti+1, "rows": rows})
    textboxes = extract_docx_textbox_structure(filepath)
    if textboxes:
        write_log(f"INFO: 检测到 {len(textboxes)} 个 Word 文本框")
    result = {"file_name": filepath.name, "file_type": "docx", "paragraphs": paras, "tables": tables}
    if textboxes:
        result["textboxes"] = textboxes
    return result

def extract_xlsx_form_structure(filepath: Path) -> dict:
    """提取 xlsx 空白表结构供 API 使用。"""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
    except Exception as e:
        return {"file_name": filepath.name, "file_type": "xlsx", "error": str(e)}
    sheets = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [{"col_index": ci+1, "text": str(c).strip() if c is not None else ""} for ci, c in enumerate(row)]
            rows.append({"row_index": len(rows)+1, "cells": cells})
        sheets.append({"sheet_name": sn, "rows": rows})
    wb.close()
    return {"file_name": filepath.name, "file_type": "xlsx", "sheets": sheets}


# ------------------------------------------------------------
# Word 文本框写入
# ------------------------------------------------------------
def _is_textbox_safe_to_overwrite(text: str, label: str = "") -> bool:
    """
    Check if a textbox is safe to overwrite.
    Safe: empty, or only contains label (e.g. "个人简介："), or label+placeholder (e.g. "个人简介：____").
    Not safe: already has substantial content.
    """
    stripped = text.strip()
    if not stripped:
        return True
    # Only label or label + colon
    label_variants = [label, label + "：", label + ":", label + "：____", label + ":____",
                      label + "：________", label + ":________", label + "：  ", label + ":  "]
    for v in label_variants:
        if stripped == v:
            return True
    # Label followed by only whitespace/punctuation
    if label and stripped.startswith(label):
        rest = stripped[len(label):].strip().strip("：:").strip("_").strip()
        if not rest:
            return True
    # Very short text that looks like just a label (< 20 chars)
    if len(stripped) < 20 and label and label in stripped:
        return True
    return False


def write_docx_textbox(draft_path, textbox_index: int, value: str, label: str = "") -> tuple:
    """
    Write content into a specific textbox in a .docx draft file.
    Uses zipfile + xml.etree.ElementTree to modify word/document.xml.
    Returns (success: bool, message: str).
    """
    import tempfile
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    try:
        # Read original
        with zipfile.ZipFile(str(draft_path), "r") as zf:
            doc_xml = zf.read("word/document.xml")
            all_files = {name: zf.read(name) for name in zf.namelist()}
    except Exception as e:
        return False, f"无法读取草稿: {e}"

    try:
        root = ET.fromstring(doc_xml)
    except ET.ParseError as e:
        return False, f"XML 解析失败: {e}"

    # Collect all txbxContent elements
    tbxcs = []
    for el in root.iter(f"{{{W_NS}}}txbxContent"):
        tbxcs.append(el)
    # Also v:textbox
    V_NS = "urn:schemas-microsoft-com:vml"
    for el in root.iter(f"{{{V_NS}}}textbox"):
        tbxcs.append(el)
    # Also wps:txbx
    WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    for el in root.iter(f"{{{WPS_NS}}}txbx"):
        tbxcs.append(el)

    if textbox_index < 1 or textbox_index > len(tbxcs):
        return False, f"文本框索引 {textbox_index} 越界（共 {len(tbxcs)} 个文本框）"

    tb = tbxcs[textbox_index - 1]
    # Read current text
    t_nodes = list(tb.iter(f"{{{W_NS}}}t"))
    old_text = "".join(t.text or "" for t in t_nodes).strip()

    if not _is_textbox_safe_to_overwrite(old_text, label):
        return False, f"文本框 #{textbox_index} 已有真实内容（{old_text[:40]}），禁止覆盖"

    # Build new content
    if label:
        new_text = label + "：" + value
    else:
        new_text = value

    # Replace text: if we have t nodes, replace the first one and clear the rest
    if t_nodes:
        # Set first t node to the full new text
        t_nodes[0].text = new_text
        t_nodes[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        # Clear remaining t nodes
        for tn in t_nodes[1:]:
            tn.text = ""
    else:
        # No existing w:t nodes — create a paragraph with text
        p = ET.Element(f"{{{W_NS}}}p")
        r = ET.SubElement(p, f"{{{W_NS}}}r")
        t = ET.SubElement(r, f"{{{W_NS}}}t")
        t.text = new_text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        # Insert as first child of the textbox
        tb.insert(0, p)

    # Write back — rebuild the zip with modified document.xml
    new_xml = ET.tostring(root, encoding="unicode", xml_declaration=True)
    # Ensure correct encoding header
    if not new_xml.startswith("<?xml"):
        new_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>' + new_xml

    tmp_path = Path(str(draft_path) + ".tmp")
    try:
        with zipfile.ZipFile(str(tmp_path), "w", zipfile.ZIP_DEFLATED) as zout:
            for name, data in all_files.items():
                if name == "word/document.xml":
                    zout.writestr(name, new_xml.encode("utf-8"))
                else:
                    zout.writestr(name, data)
    except Exception as e:
        if tmp_path.exists():
            tmp_path.unlink()
        return False, f"写入 zip 失败: {e}"

    # Replace original
    tmp_path.replace(draft_path)
    return True, f"文本框 #{textbox_index} 写入成功"


# ------------------------------------------------------------
# 填写计划归一化（兼容新旧格式）
# ------------------------------------------------------------
def normalize_fill_plan(fill_plan: dict) -> dict:
    """
    Normalize fill plan writes to internal target format.
    Compatible with:
      A. Old nested: {"target":{"type":"docx_textbox","textbox_index":1},...}
      B. New flat:   {"target_type":"docx_textbox","textbox_index":1,...}
    """
    writes = fill_plan.get("writes", [])
    normalized = []
    for w in writes:
        nw = {"field_label": w.get("field_label", ""),
              "field_key": w.get("field_key", ""),
              "value": w.get("value", ""),
              "reason": w.get("reason", "")}

        # Already has nested target -> keep as-is
        if "target" in w and isinstance(w["target"], dict):
            nw["target"] = dict(w["target"])
        # Flat format -> convert to nested target
        elif "target_type" in w:
            tgt = {"type": w["target_type"]}
            if w["target_type"] == "docx_table":
                tgt["table_index"] = w.get("table_index", 1)
                tgt["row"] = w.get("row", 1)
                tgt["col"] = w.get("col", 1)
            elif w["target_type"] == "docx_textbox":
                tgt["textbox_index"] = w.get("textbox_index", 1)
            else:
                tgt["_original_type"] = w["target_type"]
            nw["target"] = tgt
        # No target info -> still include but mark
        else:
            nw["target"] = {"type": "unknown"}

        normalized.append(nw)

    result = dict(fill_plan)
    result["writes"] = normalized
    return result


# ------------------------------------------------------------
# API-first 填写计划应用
# ------------------------------------------------------------
def _apply_fill_plan_to_draft(draft_path: Path, fill_plan: dict, profile: dict, ext: str, source_file: Path) -> dict:
    """将 API 返回的填写计划应用到草稿副本。"""
    fill_plan = normalize_fill_plan(fill_plan)
    report = {"filled": [], "skipped": [], "unknown": [], "blocked": [], "review_needed": [], "not_supported": []}
    writes = fill_plan.get("writes", [])
    if ext == ".docx":
        import docx
        try:
            doc = docx.Document(str(draft_path))
        except Exception as e:
            report["not_supported"].append(f"无法打开草稿: {e}")
            return report
        cells_filled = 0
        textbox_writes = []
        for w in writes:
            tgt = w.get("target", {})

            # Collect textbox writes (defer to after doc.save for tables)
            if tgt.get("type") == "docx_textbox":
                textbox_writes.append(w)
                continue

            # Table write path (original logic)
            ti = tgt.get("table_index", 1) - 1
            ri = tgt.get("row", 1) - 1
            ci = tgt.get("col", 1) - 1
            try:
                target_cell = doc.tables[ti].rows[ri].cells[ci]
                if not is_empty_cell(target_cell.text):
                    report["blocked"].append({"field_label": w.get("field_label",""), "reason": "已有内容"})
                    continue
                for para in target_cell.paragraphs:
                    for run in para.runs: run.text = ""
                if target_cell.paragraphs:
                    target_cell.paragraphs[0].add_run(str(w.get("value","")))
                cells_filled += 1
                report["filled"].append({"field_label": w.get("field_label",""), "field_key": w.get("field_key",""), "table": ti+1, "row": ri+1, "col": ci+1, "location": f"Table[{ti+1}] R{ri+1}C{ci+1}"})
            except (IndexError, Exception):
                report["not_supported"].append({"field_label": w.get("field_label",""), "reason": "越界"})

        # Save doc for table writes first
        doc.save(str(draft_path))

        # Now apply textbox writes to the saved zip
        for w in textbox_writes:
            tb_tgt = w.get("target", {})
            label = w.get("field_label", "")
            value = str(w.get("value", ""))
            tb_idx = tb_tgt.get("textbox_index", 1)
            ok, msg = write_docx_textbox(draft_path, tb_idx, value, label)
            if ok:
                report["filled"].append({
                    "field_label": label,
                    "field_key": w.get("field_key", ""),
                    "target_type": "docx_textbox",
                    "textbox_index": tb_idx,
                    "location": f"Textbox[{tb_idx}]",
                })
            else:
                report["not_supported"].append({
                    "field_label": label,
                    "reason": msg,
                    "target_type": "docx_textbox",
                    "textbox_index": tb_idx,
                })

        report["_stats"] = {"cells_filled": cells_filled}
    else:
        import openpyxl
        try:
            wb = openpyxl.load_workbook(str(draft_path))
        except Exception as e:
            report["not_supported"].append(f"无法打开草稿: {e}")
            return report
        cells_filled = 0
        for w in writes:
            tgt = w.get("target", {})
            try:
                ws = wb[tgt.get("sheet_name", wb.sheetnames[0])]
                cell = ws.cell(row=tgt.get("row",1), column=tgt.get("col",1))
                if not is_empty_cell(cell.value):
                    report["blocked"].append({"field_label": w.get("field_label",""), "reason": "已有内容"})
                    continue
                cell.value = str(w.get("value",""))
                cells_filled += 1
                report["filled"].append({"field_label": w.get("field_label",""), "field_key": w.get("field_key",""), "table": tgt.get("sheet_name", ws.title), "row": tgt.get("row",1), "col": tgt.get("col",1), "location": f"Sheet[{tgt.get('sheet_name','')}] R{tgt.get('row',1)}C{tgt.get('col',1)}"})
            except (IndexError, Exception):
                report["not_supported"].append({"field_label": w.get("field_label",""), "reason": "越界"})
        wb.save(str(draft_path))
        report["_stats"] = {"cells_filled": cells_filled}
    return report


# ------------------------------------------------------------
# 主流程
# ------------------------------------------------------------
def get_target_form_files(process_all: bool = False) -> list:
    """Get target form files from new_forms. Default: only the most recently modified one."""
    files = []
    for pattern in ("*.docx", "*.xlsx"):
        files.extend(NEW_FORMS_DIR.glob(pattern))
    # Filter temp files
    files = [p for p in files if not p.name.startswith("~$")]
    if not files:
        return []
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    if process_all:
        return files
    # Default: only the latest file
    if len(files) > 1:
        print(f"\n  检测到 {len(files)} 个待填写表，默认只处理最新文件：{files[0].name}")
        print(f"  如需处理全部，请使用 --all。\n")
    return [files[0]]


def main():
    write_log("========== 开始自动填写草稿 ==========")
    process_all = "--all" in sys.argv
    dry_list = "--dry-list" in sys.argv

    # 1. 检查资料库
    VAULT_PROFILE.parent.mkdir(parents=True, exist_ok=True)
    profile_path = VAULT_PROFILE

    if not profile_path.exists():
        write_log("ERROR: profile.json 不存在，无法继续。")
        print("\n" + "!" * 60)
        print("  错误：未找到 profile.json")
        print("")
        print("  请先完成第 1 阶段：本地资料库初始化。")
        print("  步骤：")
        print(f"    1. 把旧表格放入 {PROJECT_ROOT / 'input_forms'}")
        print(f"    2. 运行 python {PROJECT_ROOT / 'app' / 'extract_candidates.py'}")
        print(f"    3. 检查候选信息后创建 confirmed_profile.json")
        print(f"    4. 运行 python {PROJECT_ROOT / 'app' / 'save_confirmed_profile.py'}")
        print("")
        print(f"  确认 profile.json 存在后，再运行本脚本。")
        print("!" * 60 + "\n")
        sys.exit(1)

    profile = load_profile(profile_path)
    if not profile:
        write_log("ERROR: profile.json 无有效资料。")
        print("\n错误：profile.json 中没有发现已确认的资料字段。")
        print("请确认资料库中包含至少一个已确认的字段值。\n")
        sys.exit(1)

    write_log(f"INFO: 加载资料库成功，包含 {len(profile)} 个字段。")
    print(f"资料库: {profile_path.name} ({len(profile)} 个字段)")

    # 2. 检查新表格目录
    NEW_FORMS_DIR.mkdir(parents=True, exist_ok=True)
    all_new_files = get_target_form_files(process_all=process_all)

    # --dry-list: only show what would be processed, then exit
    if dry_list:
        print(f"\n  --dry-list 模式（{'全部' if process_all else '默认：最新文件'}）")
        if not all_new_files:
            print(f"  没有待处理的文件。")
        else:
            print(f"  将处理 {len(all_new_files)} 个文件：")
            for f in all_new_files:
                print(f"    - {f.name} ({datetime.fromtimestamp(f.stat().st_mtime).strftime('%Y-%m-%d %H:%M:%S')})")
        print()
        return

    if not all_new_files:
        write_log("INFO: new_forms 中没有可处理的文件。")
        print("\n[提示] new_forms 中没有发现 .docx 或 .xlsx 文件。")
        print(f"   请将待填写的新表格放入：{NEW_FORMS_DIR}")
        print(f"   支持格式：.docx（Word文档）、.xlsx（Excel表格）")
        print(f"   不支持：.pdf、扫描件、图片\n")
        return

    # 忽略非目标格式
    ignored = []
    for f in NEW_FORMS_DIR.iterdir():
        if f.is_file():
            ext = f.suffix.lower()
            if ext not in (".docx", ".xlsx"):
                ignored.append(f.name)

    if ignored:
        write_log(f"INFO: 忽略非目标格式文件: {', '.join(ignored)}")
        print(f"  忽略文件（格式不支持）: {', '.join(ignored)}")

    docx_count = sum(1 for f in all_new_files if f.suffix == ".docx")
    xlsx_count = sum(1 for f in all_new_files if f.suffix == ".xlsx")
    write_log(f"INFO: 发现 {docx_count} 个 .docx 文件，{xlsx_count} 个 .xlsx 文件{', 仅处理最新' if not process_all and len(all_new_files) > 0 else ''}。")

    # ---- API-first 检查 ----
    use_api = False
    try:
        import sys as _sys; _sys.path.insert(0, str(PROJECT_ROOT / "app"))
        from api_assist import assert_api_ready, send_json_request
        ok, msg = assert_api_ready()
        if ok:
            use_api = True
        else:
            print(f"\n  API-first 未就绪: {msg}")
            print(f"  回退到本地填写。\n")
    except ImportError:
        print("\n  api_assist 模块不可用，使用本地填写。\n")
    except Exception as e:
        print(f"\n  API 配置异常: {e}，回退本地填写。\n")

    # 3. 逐文件处理
    DRAFT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    success_count = 0
    run_entries = []
    html_entries = []

    for filepath in all_new_files:
        print(f"\n{'='*50}")
        print(f"  处理中: {filepath.name}")
        write_log(f"INFO: 开始处理: {filepath.name}")

        ext = filepath.suffix.lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        # 确定草稿文件名
        stem = filepath.stem
        draft_filename = f"{stem}_自动填写草稿{ext}"
        draft_path = DRAFT_OUTPUT_DIR / draft_filename

        # 如果已存在，加时间戳避免覆盖
        if draft_path.exists():
            draft_filename = f"{stem}_自动填写草稿_{timestamp}{ext}"
            draft_path = DRAFT_OUTPUT_DIR / draft_filename

        # 复制原始文件作为草稿基础
        try:
            shutil.copy2(str(filepath), str(draft_path))
            write_log(f"INFO: 创建草稿副本: {draft_filename}")
        except Exception as e:
            write_log(f"ERROR: 无法复制文件 {filepath.name}: {e}")
            print(f"  [X] 复制失败: {e}")
            continue

        # 根据类型填写
        try:
            if use_api:
                # API-first: get fill plan from LLM
                if ext == ".docx":
                    form_structure = extract_docx_form_structure(filepath)
                else:
                    form_structure = extract_xlsx_form_structure(filepath)
                payload = {
                    "task": "form_fill_plan",
                    "instruction": (
                        "请根据 vault 个人资料和空白表结构生成填写计划。只填写明确对应的个人信息字段，不要编造信息。"
                        "使用扁平 JSON 格式，不要嵌套 target 对象。"
                        "每个 write 必须包含: target_type, table_index, row, col, textbox_index, field_label, field_key, value。"
                        "表格写入: target_type=docx_table, 填 table_index/row/col, textbox_index=null。"
                        "文本框写入: target_type=docx_textbox, 填 textbox_index, table_index/row/col=null。"
                        "文本框 label_guess 映射: 个人简介→biography, 项目经历/项目经历及本人贡献→project_experience, 研究方向→research_area。"
                        "空文本框或仅含标签的可以写入。已有内容的不要覆盖，放入 need_user_review。"
                        "只输出 JSON，不要 Markdown 或解释文字。"
                    ),
                    "rules": [
                        "不要覆盖已有内容",
                        "不要填写课程表字段",
                        "不确定放入 unknown_fields",
                        "不要忽略 textboxes",
                        "不要嵌套 target 对象，使用扁平字段",
                        "只输出纯 JSON",
                    ],
                    "vault_profile": {k: str(v)[:200] for k, v in profile.items()},
                    "form_structure": form_structure,
                }
                schema_hint = (
                    '{"writes":['
                    '{"target_type":"docx_table","table_index":1,"row":1,"col":2,"textbox_index":null,'
                    '"field_label":"姓名","field_key":"name","value":"...","reason":""},'
                    '{"target_type":"docx_textbox","table_index":null,"row":null,"col":null,"textbox_index":1,'
                    '"field_label":"个人简介","field_key":"biography","value":"...","reason":""}'
                    '],"unknown_fields":[],"need_user_review":[]}'
                )
                api_ok, fill_plan, api_err = send_json_request("form_fill_plan", payload, schema_hint)
                if api_ok and fill_plan:
                    report = _apply_fill_plan_to_draft(draft_path, fill_plan, profile, ext, filepath)
                    # Note textboxes for user review
                    if form_structure.get("textboxes"):
                        tb_count = len(form_structure["textboxes"])
                        report["review_needed"].append({
                            "field_label": "Word 文本框",
                            "reason": f"检测到 {tb_count} 个文本框，当前仅识别结构暂不自动写入，请人工检查",
                            "textboxes": [{"index": tb["textbox_index"], "label_guess": tb["label_guess"]} for tb in form_structure["textboxes"]],
                        })
                    write_log(f"INFO: API 填写计划已应用 ({len(fill_plan.get('writes',[]))} writes)")
                    print(f"  API 填写: {len(fill_plan.get('writes',[]))} writes")
                else:
                    write_log(f"WARN: API 填写计划失败: {api_err}，回退本地")
                    if ext == ".docx":
                        report = fill_docx(draft_path, profile)
                    else:
                        report = fill_xlsx(draft_path, profile)
                    if form_structure.get("textboxes"):
                        report["review_needed"].append({
                            "field_label": "Word 文本框",
                            "reason": f"检测到 {len(form_structure['textboxes'])} 个文本框，当前仅识别结构暂不自动写入，请人工检查",
                        })
            else:
                if ext == ".docx":
                    report = fill_docx(draft_path, profile)
                else:
                    report = fill_xlsx(draft_path, profile)
        except Exception as e:
            write_log(f"ERROR: 处理 {filepath.name} 时异常: {e}")
            print(f"  [X] 处理异常: {e}")
            continue

        # 生成报告
        json_path, md_path = generate_reports(
            filepath, draft_path, profile_path, report, profile, timestamp
        )

        # 打印摘要
        filled_count = len(report.get("filled", []))
        review_count = len(report.get("review_needed", []))
        missed_count = len(report.get("skipped", [])) + len(report.get("unknown", []))
        blocked_count = len(report.get("blocked", []))

        print(f"  [OK] 草稿: {draft_filename}")
        print(f"  [STATS] 已填写: {filled_count} | 需复核: {review_count} | 未填写: {missed_count} | 已阻塞: {blocked_count}")
        print(f"  [DOC] 报告: {md_path.name}")

        # 生成 HTML 预览
        try:
            html_path = generate_html_preview(filepath, draft_path, profile, report, timestamp)
            print(f"  [HTML] 预览: {html_path.name}")
        except Exception as e:
            write_log(f"WARN: HTML 预览生成失败 {filepath.name}: {e}")
            html_path = None

        success_count += 1
        run_entries.append({
            "source_file": str(filepath),
            "draft_file": str(draft_path),
            "report_json": str(json_path),
            "report_md": str(md_path),
            "filled_count": filled_count,
        })
        if html_path:
            html_entries.append({
                "source_file": str(filepath),
                "draft_file": str(draft_path),
                "html_file": str(html_path),
                "html_json": str(HTML_OUTPUT_DIR / f"{re.sub(r'[\\\\/:*?\"<>|]', '_', filepath.stem)[:40]}_{timestamp}.json"),
                "filled_count": filled_count,
            })

    # 写入 latest_filling_run.json
    if run_entries:
        latest = {
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "reports": run_entries,
        }
        REPORT_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        latest_path = REPORT_OUTPUT_DIR / "latest_filling_run.json"
        with open(latest_path, "w", encoding="utf-8") as f:
            json.dump(latest, f, ensure_ascii=False, indent=2)
        write_log(f"INFO: 已生成 latest_filling_run.json ({len(run_entries)} 个报告)")

    # 写入 latest_review_html.json
    if html_entries:
        write_latest_review_html(html_entries)
        write_log(f"INFO: 已生成 latest_review_html.json ({len(html_entries)} 个 HTML 预览)")

    # 4. 总结
    print(f"\n{'='*60}")
    print(f"  第 2 阶段完成！")
    print(f"  处理文件: {len(all_new_files)} 个")
    print(f"  成功生成草稿: {success_count} 个")
    print(f"  草稿目录: {DRAFT_OUTPUT_DIR}")
    print(f"  报告目录: {REPORT_OUTPUT_DIR}")
    if html_entries:
        print(f"  HTML 预览: {HTML_OUTPUT_DIR}")
    print(f"")
    print(f"  [!] 重要提醒:")
    print(f"  - 草稿不是最终版，请在第 3 阶段检查")
    print(f"  - 本阶段未联网、未调用 API")
    print(f"  - 原始文件未被修改")
    print(f"  - 资料库未被修改")
    print(f"  - 未自动保存任何新信息")
    print(f"{'='*60}\n")

    write_log(f"========== 自动填写草稿完成（成功 {success_count}/{len(all_new_files)}） ==========")


if __name__ == "__main__":
    main()
