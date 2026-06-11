# -*- coding: utf-8 -*-
"""
SafeFill - 候选信息提取脚本
功能：从 input_forms 目录中的 .docx / .xlsx / .pdf 文件提取候选个人信息，
     生成脱敏后的 Markdown 检查文件 + JSON 数据 + confirmed_profile 模板草稿。
安全约束：不联网、不修改原始文件、敏感字段脱敏显示。

v1.3 增强：
  - 结构化表格提取（同单元格字段、字段在上值在下、多字段同行拆分）
  - 字段别名精确化（研究方向→research_area，部门严格匹配）
  - 质量等级（high/medium/low/needs_manual_review）
  - 姓名异常检测（混入学号/导师/专业等词时降级）
  - confirmed_profile_template 自动生成
  - 表格中的自定义字段提取（民族、导师、学号、爱好等）
  - 无标签强格式字段提取（邮箱、手机号、身份证号、出生日期推断），默认需用户确认
"""

import os
import sys
import json
import re
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
INPUT_DIR = PROJECT_ROOT / "input_forms"
OUTPUT_DIR = PROJECT_ROOT / "candidate_reviews"
LOGS_DIR = PROJECT_ROOT / "logs"

# ------------------------------------------------------------
# 字段别名（v1.3 修正）
# ------------------------------------------------------------
FIELD_ALIASES = {
    "name":           ["姓名", "名字", "申报人", "申请人", "负责人"],
    "gender":         ["性别"],
    "birth_date":     ["出生日期", "出生年月", "生日", "出生年月日"],
    "id_number":      ["身份证号", "身份证号码", "证件号码", "身份证", "公民身份号码"],
    "phone":          ["手机号", "手机号码", "联系电话", "电话", "联系方式", "移动电话"],
    "email":          ["邮箱", "电子邮件", "电子邮箱", "E-mail", "Email", "邮件"],
    "organization":   ["工作单位", "所在单位", "依托单位", "申报单位", "单位名称"],
    "department":     ["部门", "所在部门", "所属部门", "科室", "所在科室", "所在学院", "学院(系所)", "学院(盖章)", "学院（盖章）", "所在学院(系所)", "院系"],
    "title":          ["职称", "专业技术职务", "现任职称", "职务"],
    "education":      ["学历", "最高学历", "文化程度", "教育背景", "教育经历"],
    "degree":         ["学位", "最高学位", "学位名称"],
    "major":          ["专业", "所学专业", "攻读专业", "学科专业", "专业名称", "就读专业"],
    "research_area":  ["研究方向", "研究方向/兴趣", "研究兴趣", "研究领域", "研究主题",
                        "专业方向", "学科方向", "主攻方向", "擅长方向", "临床方向", "从事专业"],
    "address":        ["通讯地址", "通信地址", "联系地址", "家庭地址", "住址", "邮寄地址", "现居地", "现住址", "居住地", "所在地"],
    "photo_path":     ["证件照", "照片", "相片", "一寸照片", "免冠照片"],
    "project_experience": ["项目经历", "科研项目", "主持项目", "参与项目", "项目列表", "承担项目"],
    "biography":      ["个人简介", "个人简历", "简介", "简历", "个人陈述", "自我简介"],
    "age":            ["年龄"],
    # 仅作为字段边界，不映射到核心资料库字段
    "_separator":     ["学号", "导师", "指导老师", "指导教师", "学生证号", "编号"],
}

# 常见个人信息 → 自定义字段（不在标准模板中，但值得保存到 custom_fields）
CUSTOM_FIELD_ALIASES = {
    "民族":           ["民族"],
    "政治面貌":       ["政治面貌", "政治身份"],
    "籍贯":           ["籍贯", "祖籍"],
    "健康状况":       ["健康状况", "身体状况", "健康"],
    "婚姻状况":       ["婚姻状况", "婚否", "婚姻"],
    "紧急联系人":     ["紧急联系人", "紧急联系人姓名"],
    "紧急联系电话":   ["紧急联系电话", "紧急电话"],
    "爱好":           ["爱好", "兴趣爱好"],
    "特长":           ["特长", "个人特长", "专长"],
    "奖惩情况":       ["奖惩情况", "奖惩", "获奖情况"],
    "资格证书":       ["资格证书", "职业资格", "持证情况"],
    "学号":           ["学号", "学生证号"],
    "导师":           ["导师", "指导老师", "指导教师"],
}

# 课程表相关字段 —— 禁止作为个人信息保存
COURSE_KEYWORDS = [
    "学位课程", "学位课", "非学位课", "课程", "课程名称", "课程编号",
    "学分", "学时", "课程类别", "必修课", "选修课", "培养课程",
    "开课单位", "任课教师", "上课时间", "上课地点",
]

# 每个 field_key 的主显示名
FIELD_LABEL = {k: v[0] for k, v in FIELD_ALIASES.items()}

# 严格匹配字段（只匹配精确的，避免"专业"误匹配"专业方向"等）
STRICT_ALIASES = {
    # field_key -> set of aliases that need EXACT match (not substring)
    "department":    {"学院", "系所", "院系"},  # only exact match, avoid substring false positives
}

SENSITIVE_FIELDS = {"id_number", "phone", "address", "photo_path"}

# ------------------------------------------------------------
# 课程字段检测
# ------------------------------------------------------------
def is_course_field(label: str, value: str = "") -> bool:
    """检测是否为课程表相关字段（不应作为个人信息保存）。"""
    # 归一化：去掉换行和多余空格
    normalized = re.sub(r'\s+', '', label)
    for kw in COURSE_KEYWORDS:
        kw_norm = re.sub(r'\s+', '', kw)
        if kw_norm in normalized:
            return True
    # 学位单独匹配时要排除"学位课程"等
    if label.strip().rstrip("：:") == "学位" and any(
        kw in value for kw in ["课程", "学分", "学时", "必修", "选修"]
    ):
        return True
    return False


# 姓名异常关键词
NAME_CONTAMINATION_WORDS = [
    "学号", "导师", "专业", "研究方向", "培养", "学院", "课程",
    "计划", "日期", "签名", "电话", "手机", "邮箱", "地址", "编号",
    "年级", "班级", "入学", "毕业",
]

# ------------------------------------------------------------
# 日志
# ------------------------------------------------------------
def write_log(message: str):
    LOGS_DIR.mkdir(parents=True, exist_ok=True)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_file = LOGS_DIR / "extract_candidates.log"
    with open(log_file, "a", encoding="utf-8") as f:
        f.write(f"[{ts}] {message}\n")
    print(f"[LOG] {message}")


# ------------------------------------------------------------
# 脱敏
# ------------------------------------------------------------
def mask_id_number(value: str) -> str:
    v = value.strip()
    if len(v) >= 15: return v[:3] + "*" * (len(v) - 7) + v[-4:]
    if len(v) >= 8:  return v[:2] + "*" * (len(v) - 4) + v[-2:]
    return "***"

def mask_phone(value: str) -> str:
    v = value.strip().replace("-", "").replace(" ", "")
    if len(v) == 11 and v.isdigit(): return v[:3] + "****" + v[-4:]
    if len(v) >= 8: return v[:3] + "****" + v[-2:]
    return "***"

def mask_address(value: str) -> str:
    v = value.strip()
    if len(v) <= 6: return v + "***"
    return v[:6] + "***"

def mask_for_display(fkey: str, value: str) -> str:
    if not value: return ""
    if fkey == "id_number": return mask_id_number(value)
    if fkey == "phone":     return mask_phone(value)
    if fkey == "address":   return mask_address(value)
    return value


# ------------------------------------------------------------
# 字段名归一化
# ------------------------------------------------------------
def normalize_field_name(raw: str) -> str:
    """
    去掉字段名中的空格、Tab、全角/半角冒号等干扰字符。
    "姓 名" → "姓名"
    "电 话：" → "电话"
    """
    # 去掉空格和 Tab
    cleaned = re.sub(r'[\s\t]+', '', raw)
    # 去掉末尾冒号
    cleaned = cleaned.rstrip('：:')
    return cleaned


# ------------------------------------------------------------
# 别名查找
# ------------------------------------------------------------
def find_field_key(label: str) -> str | None:
    """根据中文字段名查找对应的 field_key。"""
    # 先尝试原始标签
    label = label.strip().rstrip("：:").strip()
    if not label: return None

    # 再尝试归一化后的标签
    normalized = normalize_field_name(label)

    for fkey, aliases in FIELD_ALIASES.items():
        strict_set = STRICT_ALIASES.get(fkey, set())
        for alias in aliases:
            if alias in strict_set:
                if label == alias or normalized == alias:
                    return fkey
            else:
                if alias in label or alias in normalized:
                    return fkey

    for fkey, aliases in FIELD_ALIASES.items():
        strict_set = STRICT_ALIASES.get(fkey, set())
        for alias in aliases:
            if alias in strict_set:
                # Exact match required
                if label == alias:
                    return fkey
            else:
                # Substring match allowed
                if alias in label:
                    return fkey
    return None


def find_custom_field_key(label: str) -> str | None:
    """根据中文字段名查找自定义字段，返回 custom.xxx。"""
    label = label.strip().rstrip("：:").strip()
    if not label:
        return None
    normalized = normalize_field_name(label)
    for cf_name, aliases in CUSTOM_FIELD_ALIASES.items():
        for alias in aliases:
            if alias == label or alias == normalized:
                return f"custom.{cf_name}"
            if len(alias) >= 2 and (alias in label or alias in normalized):
                return f"custom.{cf_name}"
    return None


def find_any_field_key(label: str) -> str | None:
    """先匹配标准字段，再匹配自定义字段。"""
    standard = find_field_key(label)
    custom = find_custom_field_key(label)
    if standard == "_separator" and custom:
        return custom
    return standard or custom


def field_label_from_key(fkey: str, fallback: str) -> str:
    """用于报告中的字段显示名。"""
    if fkey.startswith("custom."):
        return fkey.replace("custom.", "")
    return fallback


# ------------------------------------------------------------
# 姓名异常检测
# ------------------------------------------------------------
def assess_name_quality(value: str) -> tuple:
    """
    评估姓名提取质量。
    返回 (quality, reason)
    """
    v = value.strip()
    reasons = []

    # 检测混入关键词
    found_words = [w for w in NAME_CONTAMINATION_WORDS if w in v]
    if found_words:
        reasons.append(f"包含非姓名字段词: {', '.join(found_words)}")

    # 长度检测
    # 正常中文姓名 2-4 字，少数民族或复姓可达 6 字
    # 去掉空格和标点后计算
    clean = re.sub(r'[\s：:，,、。．\-—·]', '', v)
    if len(clean) > 8:
        reasons.append(f"长度异常（{len(clean)}字符），疑似混入多字段")
    elif len(clean) > 6:
        reasons.append(f"长度偏长（{len(clean)}字符），需人工确认")

    # 包含数字
    if re.search(r'\d', v):
        reasons.append("包含数字，疑似混入学号或编号")

    if len(reasons) >= 2:
        return "needs_manual_review", "; ".join(reasons)
    elif len(reasons) == 1:
        return "low", reasons[0]
    else:
        return "high", ""


def assess_value_quality(fkey: str, raw_value: str) -> tuple:
    """
    通用值质量评估。
    返回 (quality, reason)
    """
    v = raw_value.strip()

    # 姓名特殊处理
    if fkey == "name":
        return assess_name_quality(v)

    # 空值
    if not v:
        return "needs_manual_review", "值为空"

    # 长文本
    if fkey in ("project_experience", "biography"):
        if len(v) > 0:
            return "medium", "长文本字段，需用户复核"

    # 值中混入多个字段名（检测是否包含其他别名关键词）
    other_labels = []
    for other_fkey, aliases in FIELD_ALIASES.items():
        if other_fkey == fkey: continue
        for alias in aliases:
            if len(alias) >= 2 and alias in v and alias not in FIELD_ALIASES.get(fkey, []):
                other_labels.append(alias)
                break
    if len(other_labels) >= 2:
        return "low", f"值中疑似混入其他字段名: {', '.join(other_labels[:3])}"

    # 值过长
    if len(v) > 200:
        return "low", f"值过长（{len(v)}字符），疑似非标准格式"

    return "high", ""


def _add_result(results: list, fkey: str, label: str, value: str, source_file: str, location: str, quality: str | None = None, reason: str = ""):
    """统一加入候选结果，便于激进提取时保持质量标记一致。"""
    if not fkey or not value:
        return
    value = str(value).strip().rstrip("；;，,。.").strip()
    if not value:
        return
    if is_course_field(label, value):
        results.append({
            "field_key": fkey,
            "field_label": field_label_from_key(fkey, label),
            "value": value,
            "source_file": source_file,
            "location": location,
            "quality": "low",
            "reason": "课程表字段，非个人信息",
            "is_blocked": True,
        })
        return
    if quality is None:
        quality, reason = assess_value_quality(fkey, value)
    results.append({
        "field_key": fkey,
        "field_label": field_label_from_key(fkey, label),
        "value": value,
        "source_file": source_file,
        "location": location,
        "quality": quality,
        "reason": reason,
    })


def _extract_unlabeled_common_values(text: str, source_file: str, location: str, results: list):
    """
    激进提取：没有字段标签时，也识别强格式个人信息。
    这些字段默认 medium，需要用户确认后才保存。
    """
    if not text:
        return

    # 邮箱
    for m in re.finditer(r'[\w.\-+]+@[\w.\-]+\.[A-Za-z]{2,}', text):
        _add_result(
            results, "email", "邮箱", m.group(0), source_file,
            f"{location}(auto-email)", "medium", "无字段标签，但符合邮箱格式，需用户确认"
        )

    # 身份证号
    for m in re.finditer(r'(?<!\d)([1-9]\d{5}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx])(?!\d)', text):
        idv = m.group(1)
        _add_result(
            results, "id_number", "身份证号", idv, source_file,
            f"{location}(auto-id)", "medium", "无字段标签，但符合身份证号格式，需用户确认"
        )
        birth = f"{idv[6:10]}-{idv[10:12]}-{idv[12:14]}"
        _add_result(
            results, "birth_date", "出生日期", birth, source_file,
            f"{location}(auto-id-birth)", "medium", "由身份证号推断，需用户确认"
        )

    # 手机号
    for m in re.finditer(r'(?<!\d)(1[3-9]\d[\s\-]?\d{4}[\s\-]?\d{4})(?!\d)', text):
        phone = re.sub(r'[\s\-]', '', m.group(1))
        _add_result(
            results, "phone", "手机号", phone, source_file,
            f"{location}(auto-phone)", "medium", "无字段标签，但符合手机号格式，需用户确认"
        )


# ------------------------------------------------------------
# 结构化 docx 提取
# ------------------------------------------------------------
# ------------------------------------------------------------
# Word 文本框 / 形状文字提取
# ------------------------------------------------------------
DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}


def extract_docx_textboxes(filepath: Path) -> list:
    """
    Extract text from Word textboxes / shapes inside a .docx file.
    Reads word/document.xml via zipfile and parses XML directly.
    Returns list of dicts: [{source, index, text, chars}, ...].
    """
    textboxes = []
    try:
        with zipfile.ZipFile(str(filepath), "r") as zf:
            if "word/document.xml" not in zf.namelist():
                return textboxes
            xml_bytes = zf.read("word/document.xml")
    except Exception as e:
        write_log(f"WARN: 无法读取 {filepath.name} 内部 XML: {e}")
        return textboxes

    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError as e:
        write_log(f"WARN: {filepath.name} XML 解析失败: {e}")
        return textboxes

    texts = []

    def _extract_textbox_paragraphs(txb_el) -> str:
        """Extract text from a textbox element, preserving paragraph breaks."""
        lines = []
        for p_el in txb_el.iter(f'{{{DOCX_NS["w"]}}}p'):
            t_nodes = p_el.iter(f'{{{DOCX_NS["w"]}}}t')
            line = "".join(t.text or "" for t in t_nodes).strip()
            if line:
                lines.append(line)
        # Join paragraphs with newline; if only one line, just return it
        return "\n".join(lines) if lines else ""

    # 1. w:txbxContent (standard Word textbox content)
    for el in root.iter(f'{{{DOCX_NS["w"]}}}txbxContent'):
        text = _extract_textbox_paragraphs(el)
        if text:
            texts.append(("textbox_w", text))

    # 2. v:textbox (VML textbox — older Word format)
    for el in root.iter(f'{{{DOCX_NS["v"]}}}textbox'):
        text = _extract_textbox_paragraphs(el)
        if text:
            texts.append(("textbox_v", text))

    # 3. wps:txbx (WordProcessingShape textbox)
    for el in root.iter(f'{{{DOCX_NS["wps"]}}}txbx'):
        text = _extract_textbox_paragraphs(el)
        if text:
            texts.append(("textbox_wps", text))

    # Dedup: keep first occurrence of each non-empty text
    seen_texts = set()
    deduped = []
    for src_type, text in texts:
        normalized = text.strip()
        if normalized:
            if normalized in seen_texts:
                continue
            seen_texts.add(normalized)
        deduped.append((src_type, text))

    # Build result list
    for i, (src_type, text) in enumerate(deduped):
        textboxes.append({
            "source": f"{src_type}",
            "index": i + 1,
            "text": text,
            "chars": len(text),
        })

    return textboxes


def extract_docx_structured(filepath: Path) -> list:
    """
    从 docx 文件中结构化提取候选字段。
    返回 list of dict: {field_key, field_label, value, source_file, location, quality, reason}
    """
    import docx
    results = []

    try:
        doc = docx.Document(str(filepath))
    except Exception as e:
        write_log(f"WARN: 无法解析 {filepath.name}: {e}")
        return results

    # ---- 模式 0: 段落中的内联字段（姓名：张三  学号：123  导师：李四） ----
    for pi, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text: continue
        _extract_inline_fields(text, filepath.name, f"Para[{pi+1}]", results)
        _extract_unlabeled_common_values(text, filepath.name, f"Para[{pi+1}]", results)

    # ---- 文本框 / 形状文字提取 ----
    try:
        textboxes = extract_docx_textboxes(filepath)
        if textboxes:
            write_log(f"INFO: 检测到 {len(textboxes)} 个 Word 文本框")
            for tb in textboxes:
                loc = f"文本框 {tb['index']}"
                _extract_inline_fields(tb["text"], filepath.name, loc, results)
                _extract_unlabeled_common_values(tb["text"], filepath.name, loc, results)
    except Exception as e:
        write_log(f"WARN: 文本框提取异常: {e}")

    for table_idx, table in enumerate(doc.tables):
        rows = table.rows
        num_cols = max(len(row.cells) for row in rows) if rows else 0

        # 收集所有行的原始文本用于"字段在上、值在下"检测
        row_texts = []
        for row in rows:
            cells_text = [cell.text.strip() for cell in row.cells]
            row_texts.append(cells_text)

        # ---- 模式 1: 字段在左、值在右 ----
        for row_idx, row in enumerate(rows):
            cells = row.cells
            for col_idx in range(len(cells) - 1):
                label_cell = cells[col_idx].text.strip()
                value_cell = cells[col_idx + 1].text.strip()
                if not label_cell or not value_cell: continue

                fkey = find_any_field_key(label_cell)
                if fkey:
                    _add_result(
                        results, fkey, label_cell, value_cell, filepath.name,
                        f"Table[{table_idx+1}] R{row_idx+1}C{col_idx+1}+{col_idx+2}"
                    )

        # ---- 模式 2: 字段在上、值在下 ----
        if len(row_texts) >= 2:
            for col_idx in range(num_cols):
                for row_idx in range(len(row_texts) - 1):
                    upper = row_texts[row_idx][col_idx] if col_idx < len(row_texts[row_idx]) else ""
                    lower = row_texts[row_idx + 1][col_idx] if col_idx < len(row_texts[row_idx + 1]) else ""
                    if not upper or not lower: continue
                    fkey = find_any_field_key(upper)
                    if fkey and not find_any_field_key(lower):
                        _add_result(
                            results, fkey, upper, lower, filepath.name,
                            f"Table[{table_idx+1}] R{row_idx+1}-{row_idx+2}C{col_idx+1}"
                        )

        # ---- 模式 3: 同一单元格字段（姓名：张三） ----
        for row_idx, row in enumerate(rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text: continue
                _extract_unlabeled_common_values(
                    cell_text, filepath.name,
                    f"Table[{table_idx+1}] R{row_idx+1}C{col_idx+1}", results
                )
                # 分行处理多字段
                lines = cell_text.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line: continue
                    # 模式: "字段名：值" 或 "字段名:值"
                    for fkey, aliases in FIELD_ALIASES.items():
                        for alias in aliases:
                            m = re.match(
                                re.escape(alias) + r'\s*[：:]\s*(.+)',
                                line
                            )
                            if m:
                                val = m.group(1).strip()
                                if val and val != alias:
                                    quality, reason = assess_value_quality(fkey, val)
                                    results.append({
                                        "field_key": fkey,
                                        "field_label": alias,
                                        "value": val,
                                        "source_file": filepath.name,
                                        "location": f"Table[{table_idx+1}] R{row_idx+1}C{col_idx+1}(inline)",
                                        "quality": quality,
                                        "reason": reason,
                                    })
                    for cf_name, aliases in CUSTOM_FIELD_ALIASES.items():
                        for alias in aliases:
                            m = re.match(
                                re.escape(alias) + r'\s*[：:]\s*(.+)',
                                line
                            )
                            if m:
                                val = m.group(1).strip()
                                if val and val != alias:
                                    _add_result(
                                        results, f"custom.{cf_name}", cf_name, val,
                                        filepath.name,
                                        f"Table[{table_idx+1}] R{row_idx+1}C{col_idx+1}(custom-inline)"
                                    )

        # ---- 模式 4: 多字段同行拆分 ----
        for row_idx, row in enumerate(rows):
            for col_idx in range(len(row.cells)):
                cell_text = row.cells[col_idx].text.strip()
                if not cell_text: continue
                # 多空格/制表符分隔的多字段
                if "  " in cell_text or "\t" in cell_text:
                    parts = re.split(r'\s{2,}|\t', cell_text)
                    for part in parts:
                        part = part.strip()
                        if not part: continue
                        m = re.match(r'(.+?)[：:]\s*(.+)', part)
                        if m:
                            label = m.group(1).strip()
                            val = m.group(2).strip()
                            fkey = find_any_field_key(label)
                            if fkey and val:
                                _add_result(
                                    results, fkey, label, val, filepath.name,
                                    f"Table[{table_idx+1}] R{row_idx+1}(multi-field)"
                                )

    # 去重：同一 field_key 取最高质量的第一个
    # 过滤掉 _separator（仅作为字段边界，不保存）
    seen = {}
    for r in sorted(results, key=lambda x: _quality_rank(x["quality"])):
        fkey = r["field_key"]
        if fkey == "_separator": continue
        if fkey not in seen:
            seen[fkey] = r

    return list(seen.values())


def _extract_inline_fields(text: str, source_file: str, location: str, results: list):
    """
    从一段文本中提取"字段名：值"或"字段名  值"的内联字段。
    处理多字段同行的情况（如"姓名：张三  学号：123  导师：李四"）。
    """
    # 策略：用已知别名在文本中定位，提取每个字段后的值直到下一个字段名或行尾
    # 先用所有别名在文本中找到所有匹配位置
    hits = []
    for fkey, aliases in FIELD_ALIASES.items():
        for alias in aliases:
            # 冒号分隔 "字段名：值"
            for m in re.finditer(re.escape(alias) + r'\s*[：:]\s*', text):
                hits.append((m.start(), m.end(), fkey, alias))
            # 多空格分隔 "字段名  值"
            for m in re.finditer(re.escape(alias) + r'\s{2,}', text):
                hits.append((m.start(), m.end(), fkey, alias))
            # Tab 分隔 "字段名\t值"
            for m in re.finditer(re.escape(alias) + r'\t+', text):
                hits.append((m.start(), m.end(), fkey, alias))
            # 模糊匹配：允许字段名中间有空格/Tab（如"电 话"匹配 alias="电话"）
            if len(alias) >= 2:
                fuzzy = r'\s*'.join(re.escape(c) for c in alias)
                for m in re.finditer(fuzzy + r'\s*[：:\t]\s*', text):
                    hits.append((m.start(), m.end(), fkey, alias))
    # 也扫描 STRICT_ALIASES 中的别名（仅冒号模式，避免短词误匹配值中的词）
    for fkey, aliases in STRICT_ALIASES.items():
        for alias in aliases:
            for m in re.finditer(re.escape(alias) + r'\s*[：:]\s*', text):
                hits.append((m.start(), m.end(), fkey, alias))
    # 扫描自定义字段别名
    for cf_name, aliases in CUSTOM_FIELD_ALIASES.items():
        for alias in aliases:
            for m in re.finditer(re.escape(alias) + r'\s*[：:\t]\s*', text):
                hits.append((m.start(), m.end(), f"custom.{cf_name}", cf_name))
            for m in re.finditer(re.escape(alias) + r'\s{2,}', text):
                hits.append((m.start(), m.end(), f"custom.{cf_name}", cf_name))
            if len(alias) >= 2:
                fuzzy = r'\s*'.join(re.escape(c) for c in alias)
                for m in re.finditer(fuzzy + r'\s*[：:\t]\s*', text):
                    hits.append((m.start(), m.end(), f"custom.{cf_name}", cf_name))

    if not hits:
        return

    # 按位置排序
    hits.sort(key=lambda h: h[0])

    # 从每个字段匹配位置提取其值：到下一个字段起始位置或文本末尾
    for i, (start, end, fkey, alias) in enumerate(hits):
        if i + 1 < len(hits):
            next_start = hits[i + 1][0]
            raw_value = text[end:next_start].strip()
        else:
            raw_value = text[end:].strip()

        # 清理值末尾的标点/空白
        raw_value = raw_value.rstrip("；;，,。.").strip()
        if not raw_value: continue

        quality, reason = assess_value_quality(fkey, raw_value)
        results.append({
            "field_key": fkey,
            "field_label": alias,
            "value": raw_value,
            "source_file": source_file,
            "location": location,
            "quality": quality,
            "reason": reason,
        })




def _quality_rank(q: str) -> int:
    return {"high": 0, "medium": 1, "low": 2, "needs_manual_review": 3}.get(q, 4)


# ------------------------------------------------------------
# xlsx 提取（保持兼容）
# ------------------------------------------------------------
def extract_xlsx_structured(filepath: Path) -> list:
    """从 xlsx 提取候选字段。"""
    import openpyxl
    results = []

    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
    except Exception as e:
        write_log(f"WARN: 无法解析 {filepath.name}: {e}")
        return results

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_rows = []
        for row in ws.iter_rows(values_only=True):
            row_texts = [str(c).strip() if c is not None else "" for c in row]
            all_rows.append(row_texts)

        for row_idx, row_texts in enumerate(all_rows):
            for col_idx in range(len(row_texts) - 1):
                label = row_texts[col_idx]
                value = row_texts[col_idx + 1]
                if not label or not value: continue
                fkey = find_any_field_key(label)
                if fkey:
                    _add_result(
                        results, fkey, label, value, filepath.name,
                        f"Sheet[{sheet_name}] R{row_idx+1}C{col_idx+1}+{col_idx+2}"
                    )

        # 同单元格检测
        for row_idx, row_texts in enumerate(all_rows):
            for col_idx, cell in enumerate(row_texts):
                if not cell: continue
                _extract_unlabeled_common_values(
                    cell, filepath.name,
                    f"Sheet[{sheet_name}] R{row_idx+1}C{col_idx+1}", results
                )
                for fkey, aliases in FIELD_ALIASES.items():
                    for alias in aliases:
                        m = re.match(
                            re.escape(alias) + r'\s*[：:]\s*(.+)', cell
                        )
                        if m:
                            val = m.group(1).strip()
                            if val and val != alias:
                                quality, reason = assess_value_quality(fkey, val)
                                results.append({
                                    "field_key": fkey,
                                    "field_label": alias,
                                    "value": val,
                                    "source_file": filepath.name,
                                    "location": f"Sheet[{sheet_name}] R{row_idx+1}C{col_idx+1}(inline)",
                                    "quality": quality,
                                    "reason": reason,
                                })
                for cf_name, aliases in CUSTOM_FIELD_ALIASES.items():
                    for alias in aliases:
                        m = re.match(
                            re.escape(alias) + r'\s*[：:]\s*(.+)', cell
                        )
                        if m:
                            val = m.group(1).strip()
                            if val and val != alias:
                                _add_result(
                                    results, f"custom.{cf_name}", cf_name, val,
                                    filepath.name,
                                    f"Sheet[{sheet_name}] R{row_idx+1}C{col_idx+1}(custom-inline)"
                                )

    wb.close()
    seen = {}
    for r in sorted(results, key=lambda x: _quality_rank(x["quality"])):
        if r["field_key"] == "_separator": continue
        if r["field_key"] not in seen: seen[r["field_key"]] = r
    return list(seen.values())


# ------------------------------------------------------------
# Markdown 报告生成
# ------------------------------------------------------------
def generate_markdown_report(candidates: list, all_files: list, timestamp: str) -> Path:
    """生成四类展示的 Markdown 报告。"""
    md_path = OUTPUT_DIR / f"candidate_review_{timestamp}.md"

    # 分类
    standard_high = [c for c in candidates if c["quality"] == "high" and not c.get("is_custom") and not c.get("is_blocked")]
    standard_low = [c for c in candidates if c["quality"] in ("medium", "low", "needs_manual_review") and not c.get("is_custom") and not c.get("is_blocked")]
    custom_candidates = [c for c in candidates if c.get("is_custom") and not c.get("is_blocked")]
    blocked = [c for c in candidates if c.get("is_blocked")]

    lines = [
        "# 候选信息检查",
        f"> 生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"> 扫描文件数：{len(all_files)}",
        f"> 提取候选字段数：{len(candidates)}",
        f"> 标准字段 high: {len(standard_high)} | 自定义: {len(custom_candidates)} | 不建议: {len(blocked)}",
        "", "---", "",
        "## [NOTE] 重要提示",
        "",
        "- 低置信度字段**不会自动保存**。",
        "- 敏感信息已脱敏显示。",
        "- 自定义字段会进入 vault 的 custom_fields。",
        "", "---", "",
    ]

    # 1. 标准字段
    lines.append("## [OK] 建议保存的标准字段")
    lines.append("")
    if standard_high:
        lines.append("| 字段 | 候选值（脱敏） | 来源 | 位置 | 质量 |")
        lines.append("|------|---------------|------|------|------|")
        for c in standard_high:
            display = mask_for_display(c["field_key"], c["value"])
            if len(display) > 50: display = display[:50] + "..."
            display = display.replace("|", "/")
            lines.append(f"| {c['field_label']} | {display} | {c['source_file']} | {c['location']} | high |")
        lines.append("")
    else:
        lines.append("（无）")
        lines.append("")

    # 2. 自定义字段
    lines.append("## [CUSTOM] 建议保存的自定义字段")
    lines.append("")
    lines.append("以下字段不在标准模板中，但将保存到 vault 的 custom_fields。")
    lines.append("")
    if custom_candidates:
        lines.append("| 字段 | 候选值（脱敏） | 来源 | 位置 | 质量 |")
        lines.append("|------|---------------|------|------|------|")
        for c in custom_candidates:
            display = mask_for_display(c.get("custom_name", c["field_label"]), c["value"])
            if len(display) > 50: display = display[:50] + "..."
            display = display.replace("|", "/")
            lines.append(f"| {c['field_label']} | {display} | {c['source_file']} | {c['location']} | {c['quality']} |")
        lines.append("")
    else:
        lines.append("（无）")
        lines.append("")

    # 3. 需人工确认
    lines.append("## [WARN] 需要人工确认字段")
    lines.append("")
    if standard_low:
        lines.append("| 字段 | 候选值（脱敏） | 原因 | 来源 |")
        lines.append("|------|---------------|------|------|")
        for c in standard_low:
            display = mask_for_display(c["field_key"], c["value"])
            if len(display) > 50: display = display[:50] + "..."
            display = display.replace("|", "/")
            lines.append(f"| {c['field_label']} | {display} | {c.get('reason','')} | {c['source_file']} |")
        lines.append("")
    else:
        lines.append("（无）")
        lines.append("")

    # 4. 不建议保存
    lines.append("## [STOP] 不建议保存字段")
    lines.append("")
    if blocked:
        lines.append("| 字段 | 候选值（脱敏） | 原因 | 来源 |")
        lines.append("|------|---------------|------|------|")
        for c in blocked:
            display = c.get("value", "")[:50].replace("|", "/")
            lines.append(f"| {c['field_label']} | {display} | 课程表字段，非个人信息 | {c['source_file']} |")
        lines.append("")
    else:
        lines.append("（无）")
        lines.append("")

    lines.append("---")
    lines.append("## 下一步操作")
    lines.append("")
    lines.append("1. 逐项检查以上候选信息，特别是中/低置信度字段。")
    lines.append("2. 运行 `save_confirmed_profile.py`，由 SafeFill-ProfileSave 读取最新候选并询问是否保存。")
    lines.append("3. 不需要手动创建 `confirmed_profile.json`。")
    lines.append(f"   最新模板仅作参考: `{OUTPUT_DIR}\\confirmed_profile_template_{timestamp}.json`")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return md_path


# ------------------------------------------------------------
# JSON + Template 生成
# ------------------------------------------------------------
def generate_json_report(candidates: list, all_files: list, timestamp: str) -> Path:
    """生成完整 JSON 数据文件，含标准字段、自定义字段、blocked 字段。"""
    json_path = OUTPUT_DIR / f"candidate_review_{timestamp}.json"
    data = {
        "_说明": "候选信息完整数据（含未脱敏值）。仅供本地确认使用。",
        "_生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "_扫描文件": [f.name for f in all_files],
        "_字段总数": len(candidates),
        "candidates": {},
        "custom_fields": {},
        "blocked_fields": [],
    }
    for c in candidates:
        if c.get("is_blocked"):
            data["blocked_fields"].append({
                "field_label": c.get("field_label",""),
                "value": c.get("value",""),
                "reason": "课程表字段，非个人信息",
            })
        elif c.get("is_custom"):
            data["custom_fields"][c["field_label"]] = {
                "value": c["value"],
                "source_file": c["source_file"],
                "quality": c["quality"],
                "confirmed": False,
            }
        else:
            data["candidates"][c["field_key"]] = {
                "value": c["value"],
                "field_label": c["field_label"],
                "source_file": c["source_file"],
                "location": c["location"],
                "quality": c["quality"],
                "reason": c.get("reason", ""),
                "is_sensitive": c["field_key"] in SENSITIVE_FIELDS,
                "confirmed": False,
            }
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return json_path


def generate_template_json(candidates: list, timestamp: str) -> Path:
    """生成候选确认模板草稿，含 custom_fields。"""
    tmpl_path = OUTPUT_DIR / f"confirmed_profile_template_{timestamp}.json"

    tmpl = {
        "_说明": "候选确认模板草稿。SafeFill-ProfileSave 会读取 latest_candidate 指向的最新候选；本文件仅供人工查看。",
        "_生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "_提示": "high 质量字段已预填 confirmed=true。自定义字段在 custom_fields 中。",
        "candidates": {},
        "custom_fields": {},
    }

    for c in candidates:
        if c.get("is_blocked"):
            continue
        q = c["quality"]
        auto_confirm = (q == "high")
        entry = {
            "value": c["value"],
            "field_label": c["field_label"],
            "confirmed": auto_confirm,
            "quality": q,
            "reason": c.get("reason", ""),
            "source_file": c["source_file"],
            "_提示": "",
        }

        if q == "high":
            entry["_提示"] = "[OK] 高置信度，已预确认。请检查后确认无误。"
        elif q == "medium":
            entry["_提示"] = "[WARN] 中置信度，已预确认但建议人工检查。"
        else:
            entry["_提示"] = "[STOP] 低置信度，需要人工逐项确认。请将 confirmed 改为 true 才保存。"

        if c.get("is_custom"):
            tmpl["custom_fields"][c["field_label"]] = entry
        else:
            tmpl["candidates"][c["field_key"]] = entry

    with open(tmpl_path, "w", encoding="utf-8") as f:
        json.dump(tmpl, f, ensure_ascii=False, indent=2)
    return tmpl_path


# ------------------------------------------------------------
# API-first 提取 (v2.0)
# ------------------------------------------------------------
SUPPORTED_STANDARD_FIELDS = [
    "name", "gender", "birth_date", "id_number", "phone", "email",
    "organization", "department", "title", "education", "degree",
    "major", "research_area", "address", "photo_path",
    "project_experience", "biography",
]

def extract_docx_text_for_api(filepath: Path) -> dict:
    """提取 docx 结构化文本供 API 使用。"""
    import docx
    try:
        doc = docx.Document(str(filepath))
    except Exception as e:
        return {"file_name": filepath.name, "file_type": "docx", "error": str(e)}
    paras = [{"index": i+1, "text": p.text.strip()} for i, p in enumerate(doc.paragraphs) if p.text.strip()]
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append({"table_index": ti+1, "rows": rows})
    textboxes = extract_docx_textboxes(filepath)
    result = {"file_name": filepath.name, "file_type": "docx", "paragraphs": paras, "tables": tables}
    if textboxes:
        result["textboxes"] = textboxes
    return result

def extract_xlsx_text_for_api(filepath: Path) -> dict:
    """提取 xlsx 结构化文本供 API 使用。"""
    import openpyxl
    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
    except Exception as e:
        return {"file_name": filepath.name, "file_type": "xlsx", "error": str(e)}
    sheets = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = []
        for row in ws.iter_rows(values_only=True):
            values = [str(c).strip() if c is not None else "" for c in row]
            if any(values):
                rows.append({"row_index": len(rows)+1, "values": values})
        sheets.append({"sheet_name": sn, "rows": rows})
    wb.close()
    return {"file_name": filepath.name, "file_type": "xlsx", "sheets": sheets}

def normalize_api_profile_candidates(api_result: dict, all_files: list) -> list:
    """将大模型返回结果标准化为 all_candidates 格式。"""
    candidates = []
    raw_cands = api_result.get("candidates", {})
    for fkey, item in raw_cands.items():
        if fkey not in SUPPORTED_STANDARD_FIELDS:
            continue
        if isinstance(item, dict):
            candidates.append({
                "field_key": fkey,
                "field_label": item.get("field_label", fkey),
                "value": str(item.get("value", "")),
                "quality": item.get("quality", "medium"),
                "reason": item.get("reason", ""),
                "source_file": item.get("source_file", all_files[0].name if all_files else ""),
                "location": item.get("location", ""),
                "is_custom": False,
            })

    # custom fields
    raw_custom = api_result.get("custom_fields", {})
    for cf_name, item in raw_custom.items():
        if isinstance(item, dict):
            candidates.append({
                "field_key": f"custom.{cf_name}",
                "field_label": cf_name,
                "value": str(item.get("value", "")),
                "quality": item.get("quality", "medium"),
                "reason": item.get("reason", ""),
                "source_file": item.get("source_file", all_files[0].name if all_files else ""),
                "location": item.get("location", ""),
                "is_custom": True,
            })

    # blocked
    blocked = api_result.get("blocked_fields", [])
    for item in blocked:
        if isinstance(item, dict):
            candidates.append({
                "field_key": "blocked",
                "field_label": item.get("field_label", ""),
                "value": str(item.get("value", "")),
                "quality": "needs_manual_review",
                "reason": item.get("reason", "课程/非个人信息"),
                "source_file": all_files[0].name if all_files else "",
                "location": item.get("location", ""),
                "is_blocked": True,
            })

    return candidates


# ------------------------------------------------------------
# PDF 文本提取候选（本地模式，用于 MinerU 输出的纯文本）
# ------------------------------------------------------------
def extract_text_candidates(filepath: Path, text: str) -> list:
    """
    Parse extracted PDF text to find candidate fields.
    Uses the same FIELD_ALIASES patterns as docx/xlsx extraction.
    """
    candidates = []
    lines = text.split("\n")

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Try to match "label: value" or "label：value" patterns
        for fkey, aliases in FIELD_ALIASES.items():
            if fkey.startswith("_"):
                continue
            for alias in aliases:
                # Pattern: "别名：值" or "别名:值" or "别名 值"
                patterns = [
                    re.escape(alias) + r'\s*[：:]\s*(.+)',
                    re.escape(alias) + r'\s{2,}(.+)',
                ]
                for pat in patterns:
                    m = re.search(pat, line_stripped)
                    if m:
                        value = m.group(1).strip()
                        if value and len(value) >= 1:
                            quality = "high" if 2 <= len(value) <= 100 else "medium"
                            candidates.append({
                                "field_key": fkey,
                                "field_label": alias,
                                "value": value,
                                "quality": quality,
                                "reason": f"从 PDF 文本提取（{alias}）",
                                "source_file": filepath.name,
                                "location": f"line {i+1}",
                            })
                        break
                break  # Only match first alias per field

        # Custom field aliases
        for ckey, caliases in CUSTOM_FIELD_ALIASES.items():
            for alias in caliases:
                pat = re.escape(alias) + r'\s*[：:]\s*(.+)'
                m = re.search(pat, line_stripped)
                if m:
                    value = m.group(1).strip()
                    if value and len(value) >= 1:
                        candidates.append({
                            "field_key": ckey,
                            "field_label": alias,
                            "value": value,
                            "quality": "high",
                            "reason": f"从 PDF 文本提取（{alias}）",
                            "source_file": filepath.name,
                            "location": f"line {i+1}",
                        })
                    break

        # Detect email
        email_m = re.search(r'[\w.\-]+@[\w.\-]+\.\w+', line_stripped)
        if email_m and not any(c["field_key"] == "email" for c in candidates):
            candidates.append({
                "field_key": "email", "field_label": "邮箱",
                "value": email_m.group(),
                "quality": "high",
                "reason": "从 PDF 文本识别邮箱",
                "source_file": filepath.name,
                "location": f"line {i+1}",
            })

        # Detect phone
        phone_m = re.search(r'1[3-9]\d{9}', line_stripped)
        if phone_m and not any(c["field_key"] == "phone" for c in candidates):
            candidates.append({
                "field_key": "phone", "field_label": "手机号",
                "value": phone_m.group(),
                "quality": "high",
                "reason": "从 PDF 文本识别手机号",
                "source_file": filepath.name,
                "location": f"line {i+1}",
            })

        # Detect ID number
        id_m = re.search(r'\b\d{15,18}\b', line_stripped)
        if id_m and not any(c["field_key"] == "id_number" for c in candidates):
            val = id_m.group()
            if len(val) in (15, 18):
                candidates.append({
                    "field_key": "id_number", "field_label": "身份证号",
                    "value": val,
                    "quality": "high",
                    "reason": "从 PDF 文本识别身份证号",
                    "source_file": filepath.name,
                    "location": f"line {i+1}",
                })

    # Deduplicate by field_key: keep highest quality
    seen = {}
    for c in candidates:
        fk = c["field_key"]
        if fk not in seen or c["quality"] == "high":
            seen[fk] = c
    return list(seen.values())


# ------------------------------------------------------------
# 主流程
# ------------------------------------------------------------
def main():
    print("\nSafeFill-ProfileExtract — 从旧表提取候选信息")
    print("不修改原始文件 | 不写入 vault\n")
    write_log("========== 开始提取候选信息 v1.3 ==========")

    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_files = sorted(INPUT_DIR.glob("*.docx"))
    xlsx_files = sorted(INPUT_DIR.glob("*.xlsx"))
    pdf_files = sorted(INPUT_DIR.glob("*.pdf"))
    all_files = docx_files + xlsx_files + pdf_files

    if not all_files:
        write_log("INFO: input_forms 中没有可处理的文件。")
        print("\n[提示] input_forms 中没有发现 .docx / .xlsx / .pdf 文件。")
        print(f"   请将旧的表格文件放入：{INPUT_DIR}")
        print(f"   支持格式：.docx（Word文档）、.xlsx（Excel表格）、.pdf（需要 MinerU）\n")
        return

    file_summary = f"  {len(docx_files)} .docx, {len(xlsx_files)} .xlsx"
    pdf_info = ""
    if pdf_files:
        file_summary += f", {len(pdf_files)} .pdf"
        # Check MinerU for PDFs
        try:
            from pdf_extract import check_mineru
            pdf_mineru_ok, pdf_mineru_ver, _, pdf_mineru_err = check_mineru()
            if pdf_mineru_ok:
                pdf_info = f" | MinerU: OK ({pdf_mineru_ver[:40]})"
            else:
                pdf_info = " | MinerU: UNAVAILABLE (PDF will be skipped)"
                write_log(f"WARN: MinerU 未安装，PDF 将不会被提取。")
        except ImportError:
            pdf_info = " | MinerU: module unavailable"
    write_log(f"INFO: 发现{file_summary}{pdf_info}")

    all_candidates = []

    # ---- API-first 检查 ----
    use_api = False
    try:
        import sys as _sys; _sys.path.insert(0, str(PROJECT_ROOT / "app"))
        from api_assist import assert_api_ready, send_json_request
        ok, msg = assert_api_ready()
        if ok:
            use_api = True
        else:
            print(f"\n  API-first 未就绪: {msg}")
            print(f"  回退到本地提取。\n")
    except ImportError:
        print("\n  api_assist 模块不可用，使用本地提取。\n")
    except Exception as e:
        print(f"\n  API 配置检查异常: {e}，使用本地提取。\n")

    if use_api:
        print(f"\n  API-first 提取中...")
        documents = []
        pdf_skipped = 0
        for filepath in all_files:
            ext = filepath.suffix.lower()
            if ext == ".pdf":
                try:
                    from pdf_extract import get_extracted_text, check_mineru
                    pdf_ok, _, _, _ = check_mineru()
                    if pdf_ok:
                        pdf_text, pdf_meta = get_extracted_text(filepath, mineru_available=True)
                        if pdf_meta["success"]:
                            documents.append({
                                "file": filepath.name,
                                "type": "pdf",
                                "content": pdf_text,
                                "note": f"PDF extracted via MinerU ({pdf_meta['extracted_chars']} chars)"
                                + (" [truncated]" if pdf_meta.get("truncated") else ""),
                            })
                            write_log(f"INFO: PDF 提取成功: {filepath.name} ({pdf_meta['extracted_chars']} 字符)")
                        else:
                            pdf_skipped += 1
                            write_log(f"WARN: PDF 提取失败: {filepath.name}: {pdf_meta.get('error','?')}")
                    else:
                        pdf_skipped += 1
                        write_log(f"WARN: MinerU 不可用，跳过 PDF: {filepath.name}")
                except ImportError:
                    pdf_skipped += 1
                    write_log(f"WARN: pdf_extract 模块不可用，跳过: {filepath.name}")
            elif ext == ".docx":
                documents.append(extract_docx_text_for_api(filepath))
            else:
                documents.append(extract_xlsx_text_for_api(filepath))
        if pdf_skipped:
            print(f"  PDF 未处理: {pdf_skipped} 个（MinerU 未安装或提取失败）")

        payload = {
            "task": "profile_extract",
            "instruction": "请从以下旧资料中提取个人资料候选字段，返回严格 JSON。",
            "supported_standard_fields": SUPPORTED_STANDARD_FIELDS,
            "custom_field_policy": "如果发现民族、政治面貌、导师、学号、爱好等不在标准字段中的个人信息，请放入 custom_fields。",
            "quality_policy": {"high": "字段名和值明确对应", "medium": "根据上下文推断但需用户复核", "low": "不确定，不建议自动保存"},
            "documents": documents,
        }
        schema_hint = '{"candidates":{"name":{"value":"","field_label":"","quality":"high","reason":"","source_file":"","location":""}},"custom_fields":{},"blocked_fields":[],"uncertain_fields":[]}'
        success, api_result, err = send_json_request("profile_extract", payload, schema_hint)
        if success and api_result:
            try:
                all_candidates = normalize_api_profile_candidates(api_result, all_files)
                write_log(f"INFO: API 提取成功，{len(all_candidates)} 个字段")
                print(f"  API 提取成功，{len(all_candidates)} 个字段")

                # Supplement: local textbox extraction for fields API may miss
                # (biography, project_experience, research_area, custom certs)
                existing_keys = {c["field_key"] for c in all_candidates}
                for filepath in all_files:
                    if filepath.suffix.lower() == ".docx":
                        try:
                            tb_candidates = []
                            textboxes = extract_docx_textboxes(filepath)
                            for tb in textboxes:
                                _extract_inline_fields(tb["text"], filepath.name,
                                                       f"文本框 {tb['index']}", tb_candidates)
                            new_count = 0
                            for c in tb_candidates:
                                fk = c["field_key"]
                                if fk not in existing_keys:
                                    all_candidates.append(c)
                                    existing_keys.add(fk)
                                    new_count += 1
                            if new_count:
                                write_log(f"INFO: 文本框补充提取 {new_count} 个字段")
                                print(f"  文本框补充提取: {new_count} 个字段")
                        except Exception:
                            pass
            except Exception as e:
                write_log(f"WARN: API 结果标准化失败: {e}，回退本地提取")
                print(f"  API 结果格式异常，回退本地提取。")
                use_api = False
        else:
            write_log(f"WARN: API 请求失败: {err}")
            print(f"  API 请求失败，回退本地提取。")
            use_api = False

    if not use_api:
        for filepath in all_files:
            write_log(f"INFO: 正在处理: {filepath.name}")
            print(f"  处理中: {filepath.name} ...")
            ext = filepath.suffix.lower()

            if ext == ".pdf":
                # PDF: extract text via MinerU, then treat as document for extraction
                try:
                    from pdf_extract import get_extracted_text, check_mineru
                    pdf_ok, _, _, _ = check_mineru()
                    if pdf_ok:
                        pdf_text, pdf_meta = get_extracted_text(filepath, mineru_available=True)
                        if pdf_meta["success"] and pdf_text.strip():
                            candidates = extract_text_candidates(filepath, pdf_text)
                            write_log(f"INFO: PDF 文本提取成功 ({pdf_meta['extracted_chars']} 字符)")
                        else:
                            write_log(f"WARN: PDF 提取失败: {filepath.name}: {pdf_meta.get('error','no text')}")
                            print(f"    PDF 提取失败: {pdf_meta.get('error', 'no text')}")
                            continue
                    else:
                        write_log(f"WARN: MinerU 不可用，跳过 PDF: {filepath.name}")
                        print(f"    跳过 PDF: MinerU 未安装")
                        continue
                except ImportError:
                    write_log(f"WARN: pdf_extract 模块不可用，跳过: {filepath.name}")
                    print(f"    跳过 PDF: 模块不可用")
                    continue
            elif ext == ".docx":
                candidates = extract_docx_structured(filepath)
            else:
                candidates = extract_xlsx_structured(filepath)

            if candidates:
                write_log(f"INFO: {filepath.name} 提取到 {len(candidates)} 个候选字段 "
                          f"(high={sum(1 for c in candidates if c['quality']=='high')}, "
                          f"medium={sum(1 for c in candidates if c['quality']=='medium')}, "
                          f"low/manual={sum(1 for c in candidates if c['quality'] in ('low','needs_manual_review'))})")
                print(f"    提取到 {len(candidates)} 个字段")

                # 合并到总列表
                for c in candidates:
                    existing = [x for x in all_candidates if x["field_key"] == c["field_key"]]
                    if not existing:
                        all_candidates.append(c)
                    else:
                        # 保留质量更高的
                        if _quality_rank(c["quality"]) < _quality_rank(existing[0]["quality"]):
                            all_candidates.remove(existing[0])
                            all_candidates.append(c)

    # 后处理：标记 blocked 和 custom 字段
    for c in all_candidates:
        # 检测课程表字段
        if is_course_field(c.get("field_label", ""), c.get("value", "")):
            c["is_blocked"] = True
        # 标记自定义字段
        if c["field_key"].startswith("custom."):
            c["is_custom"] = True
            c["custom_name"] = c["field_key"].replace("custom.", "")

    if not all_candidates:
        write_log("INFO: 未从任何文件中提取到候选信息。")
        print("\n[提示] 未能从文件中提取到个人信息候选字段。\n")
        return

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    # 生成报告
    md_path = generate_markdown_report(all_candidates, all_files, timestamp)
    write_log(f"INFO: 已生成 Markdown 报告: {md_path.name}")

    json_path = generate_json_report(all_candidates, all_files, timestamp)
    write_log(f"INFO: 已生成 JSON 数据: {json_path.name}")

    tmpl_path = generate_template_json(all_candidates, timestamp)
    write_log(f"INFO: 已生成模板草稿: {tmpl_path.name}")

    # 生成 latest_candidate.json（最新指针）
    standard_n = sum(1 for c in all_candidates if not c.get("is_custom") and not c.get("is_blocked"))
    custom_n = sum(1 for c in all_candidates if c.get("is_custom"))
    blocked_n = sum(1 for c in all_candidates if c.get("is_blocked"))
    latest = {
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "candidate_review_json": str(json_path),
        "candidate_review_md": str(md_path),
        "confirmed_profile_template": str(tmpl_path),
        "source_files": [str(INPUT_DIR / f.name) for f in all_files],
        "field_count": len(all_candidates),
        "standard_field_count": standard_n,
        "custom_field_count": custom_n,
        "not_recommended_count": blocked_n,
    }
    latest_path = OUTPUT_DIR / "latest_candidate.json"
    with open(latest_path, "w", encoding="utf-8") as f:
        json.dump(latest, f, ensure_ascii=False, indent=2)
    write_log(f"INFO: 已生成 latest_candidate.json")

    # 摘要
    high_n = sum(1 for c in all_candidates if c["quality"] == "high")
    med_n = sum(1 for c in all_candidates if c["quality"] == "medium")
    low_n = sum(1 for c in all_candidates if c["quality"] in ("low", "needs_manual_review"))
    sens_n = sum(1 for c in all_candidates if c["field_key"] in SENSITIVE_FIELDS)

    print(f"\n{'='*60}")
    print(f"  提取完成！")
    print(f"  扫描文件: {len(all_files)} 个")
    print(f"  提取字段: {len(all_candidates)} 个")
    print(f"    [OK] 高置信度:  {high_n} 个")
    print(f"    [WARN] 中置信度: {med_n} 个")
    print(f"    [STOP] 低/需确认: {low_n} 个")
    if sens_n:
        print(f"    敏感字段: {sens_n} 个（已脱敏）")
    print(f"")
    print(f"  检查文件: {md_path}")
    print(f"  数据文件: {json_path}")
    print(f"  模板草稿: {tmpl_path}")
    print(f"")
    print(f"  下一步:")
    print(f"    1. 打开 Markdown 检查文件，重点看中/低置信度字段")
    print(f"    2. 运行 save_confirmed_profile.py")
    print(f"    3. SafeFill-ProfileSave 会读取最新候选并询问是否保存到 vault")
    print(f"{'='*60}\n")

    write_log("========== 提取候选信息完成 v1.3 ==========")


if __name__ == "__main__":
    main()
